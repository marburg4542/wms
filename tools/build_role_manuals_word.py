"""Create editable Word manuals using the visual system of FLIGHTLYNC-Manual-DPIM.pdf."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from build_role_manuals import ROLE

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output' / 'role-manuals-word'
SHOT = ROOT / 'output' / 'manual-screenshots'
LOGO = ROOT / 'public' / 'icons' / 'ICS.png'
OUT.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(0, 112, 184)
INK = RGBColor(35, 35, 35)
GRAY = RGBColor(98, 98, 98)

def font(run, size=11, bold=False, color=INK):
    run.font.name = 'Arial'
    for key in ('ascii', 'hAnsi', 'eastAsia'):
        run._element.rPr.rFonts.set(qn(f'w:{key}'), 'Arial')
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = color

def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run('Page '); font(r, 9, color=GRAY)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

def setup(doc, lang):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65); sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72); sec.right_margin = Inches(0.72)
    sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.32)
    normal = doc.styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(11)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.paragraph_format.line_spacing=1.25; normal.paragraph_format.space_after=Pt(7)
    for style, size, before, after in [('Heading 1',16,18,8),('Heading 2',13,12,6),('Heading 3',11,8,4)]:
        s=doc.styles[style]; s.font.name='Arial'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=INK
        s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial'); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    head=sec.header.paragraphs[0]; head.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=head.add_run('WMS User Manual' if lang=='en' else 'คู่มือผู้ใช้งาน WMS'); font(r,10,color=INK)
    r=head.add_run('\nRev. 1.0 - 20260722'); font(r,9,color=INK)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.LEFT
    footer.add_run().add_picture(str(LOGO), width=Inches(.72))
    r=footer.add_run('                                             iCreativeSystems Co., Ltd.\n                                             icsco.ai'); font(r,8.5,color=INK)
    page_field(sec.footer.add_paragraph())

def p(doc, text='', style=None, center=False, size=11, bold=False, color=INK, before=0, after=None):
    para=doc.add_paragraph(style=style)
    para.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before=Pt(before)
    if after is not None: para.paragraph_format.space_after=Pt(after)
    r=para.add_run(text); font(r,size,bold,color)
    return para

def add_bullets(doc, values):
    for value in values:
        q=doc.add_paragraph(style='List Bullet'); q.paragraph_format.space_after=Pt(3)
        font(q.add_run(value), 10.5)

def add_steps(doc, values):
    for value in values:
        q=doc.add_paragraph(style='List Number'); q.paragraph_format.space_after=Pt(3)
        font(q.add_run(value), 10.5)

def figure(doc, filename, caption):
    image_path=SHOT/filename
    if not image_path.exists(): raise FileNotFoundError(image_path)
    para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before=Pt(6); para.paragraph_format.space_after=Pt(2)
    para.add_run().add_picture(str(image_path), width=Inches(5.85))
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after=Pt(10)
    font(cap.add_run(caption),9,False,GRAY)

def cover(doc, data, lang):
    p(doc,'',before=120)
    logo=doc.add_paragraph(); logo.alignment=WD_ALIGN_PARAGRAPH.CENTER
    logo.add_run().add_picture(str(LOGO),width=Inches(3.15))
    p(doc,data['title'],center=True,size=21,bold=False,after=8)
    p(doc,data['role'],center=True,size=17,bold=False,after=8)
    p(doc,data['tag'],center=True,size=12,color=GRAY,after=58)
    p(doc,'iCreativeSystems Co., Ltd.\nicsco.ai',center=True,size=9.5,color=INK)
    doc.add_page_break()

def build(role, lang, data):
    doc=Document(); setup(doc,lang); cover(doc,data,lang)
    intro_title='บทนำ' if lang=='th' else 'Introduction'
    p(doc,intro_title,style='Heading 1')
    p(doc,data['summary'])
    p(doc,'สารบัญ' if lang=='th' else 'Contents',style='Heading 2')
    headings=[title for title,_ in data['steps']]
    entries=[intro_title, ('สิ่งที่ทำได้' if lang=='th' else 'What you can do'), *headings, ('ภาพหน้าจอจากระบบจริง' if lang=='th' else 'Screens from the live system'), ('ข้อควรระวัง' if lang=='th' else 'Important notes')]
    for item in entries:
        p(doc,item, size=10.5, before=0, after=4)
    p(doc,'สิ่งที่ทำได้' if lang=='th' else 'What you can do',style='Heading 1')
    add_bullets(doc,data['allowed'])
    p(doc,'การเข้าสู่ระบบ' if lang=='th' else 'Sign in',style='Heading 1')
    signin = ['เปิดหน้า WMS แล้วกรอกชื่อผู้ใช้และรหัสผ่าน', 'กด “เข้าสู่ระบบ” และอ่านคำแนะนำสำหรับ Role ของคุณ'] if lang=='th' else ['Open WMS and enter your username and password.', 'Select “Sign in” and read the first-use guidance for your role.']
    add_steps(doc,signin); figure(doc,'login.png','ภาพ: หน้าจอเข้าสู่ระบบ WMS' if lang=='th' else 'Figure: WMS sign-in screen')
    for idx,(title,steps) in enumerate(data['steps'],1):
        p(doc,f'{idx}. {title}',style='Heading 1')
        add_steps(doc,steps)
        # Place an actual screen directly after every documented workflow.
        matched=data['shots'][min(idx-1,len(data['shots'])-1)]
        figure(doc,*matched)
    p(doc,'ภาพหน้าจอจากระบบจริง' if lang=='th' else 'Screens from the live system',style='Heading 1')
    used=set()
    for filename, caption in data['shots']:
        if filename not in used:
            figure(doc,filename,caption); used.add(filename)
    p(doc,'ข้อควรระวัง' if lang=='th' else 'Important notes',style='Heading 1')
    add_bullets(doc,data['notes'])
    p(doc,'การแก้ไขเอกสารใน Word' if lang=='th' else 'Editing this manual in Word',style='Heading 1')
    edit_note = 'เอกสารนี้เป็นไฟล์ .docx ที่แก้ไขข้อความและปรับขนาด/เปลี่ยนภาพประกอบได้โดยตรงใน Microsoft Word' if lang=='th' else 'This is an editable .docx file. Text and embedded screenshots can be edited, resized, or replaced directly in Microsoft Word.'
    p(doc,edit_note)
    filename=f'WMS_User_Manual_{role.capitalize()}_{"TH" if lang=="th" else "EN"}_Word.docx'
    doc.save(OUT/filename)

for role, langs in ROLE.items():
    for lang, data in langs.items(): build(role,lang,data)
print(f'Created {len(list(OUT.glob("*.docx")))} editable manuals in {OUT}')
