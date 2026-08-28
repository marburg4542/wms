from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'output' / 'role-manuals'
SHOTS = ROOT / 'output' / 'manual-screenshots'
OUT.mkdir(parents=True, exist_ok=True)

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)

ROLE = {
    'viewer': {
        'th': {'title': 'คู่มือผู้ใช้งาน WMS', 'role': 'Viewer (ผู้ดูข้อมูล)', 'tag': 'ดูข้อมูลคลังและรายงานอย่างปลอดภัย',
               'summary': 'ใช้สำหรับผู้ที่ต้องการตรวจสอบภาพรวมคลัง สินค้าคงเหลือ และรายงาน โดยไม่มีสิทธิ์เปลี่ยนข้อมูลหรือส่งใบเบิก',
               'allowed': ['ดูแดชบอร์ดและประวัติการทำรายการ', 'ค้นหาและกรองสินค้าคงคลัง', 'สแกน Barcode/QR เพื่อค้นหาสินค้า', 'สร้างรายงาน PDF รายวัน รายเดือน หรือรายปี', 'ปรับโปรไฟล์ เปิดแจ้งเตือน และเปลี่ยนธีม'],
               'steps': [('เข้าสู่ระบบ', ['เปิดหน้า WMS', 'กรอกชื่อผู้ใช้และรหัสผ่าน', 'กด “เข้าสู่ระบบ”']), ('ตรวจสอบสินค้าคงคลัง', ['เลือกเมนู “สินค้าคงคลัง”', 'ค้นหาด้วยชื่อหรือ SKU หรือกดปุ่มกล้องเพื่อสแกน', 'ตรวจสอบยอดคงเหลือ หมวดหมู่ และสถานะสินค้า']), ('สร้างรายงาน', ['บน Dashboard กด “นำออก PDF”', 'เลือกรายวัน รายเดือน หรือรายปี และกำหนดช่วงเวลา', 'เลือกชนิดรายการ/โครงการตามต้องการ แล้วบันทึกไฟล์'])],
               'notes': ['บัญชี Viewer ไม่มีปุ่มเพิ่มลงใบเบิก', 'สินค้าที่ปิดใช้งานจะไม่แสดงในรายการปกติ', 'กล้องสแกนใช้ได้เมื่อเปิดระบบผ่าน HTTPS หรือ localhost เท่านั้น'],
               'shots': [('login.png', 'ภาพที่ 1: หน้าเข้าสู่ระบบ'), ('viewer-inventory.png', 'ภาพที่ 2: หน้าสินค้าคงคลังของ Viewer')]},
        'en': {'title': 'WMS User Manual', 'role': 'Viewer', 'tag': 'View inventory and reports safely',
               'summary': 'For users who need to inspect warehouse status, inventory, and reports without changing data or submitting a withdrawal request.',
               'allowed': ['View the dashboard and transaction history', 'Search and filter inventory', 'Scan a Barcode/QR code to find an item', 'Create daily, monthly, or yearly PDF reports', 'Update profile, notifications, and theme'],
               'steps': [('Sign in', ['Open WMS', 'Enter your username and password', 'Select “Sign in”']), ('Review inventory', ['Open “Inventory”', 'Search by name/SKU or select the camera button to scan', 'Review balance, category, and stock status']), ('Create a report', ['On Dashboard, select “Export PDF”', 'Choose daily, monthly, or yearly and the period', 'Select transaction/project filters if needed, then save the file'])],
               'notes': ['Viewer accounts do not see an add-to-request button.', 'Inactive items are hidden from the normal list.', 'Camera scanning requires HTTPS or localhost.'],
               'shots': [('login.png', 'Figure 1: Sign-in page'), ('viewer-inventory.png', 'Figure 2: Viewer inventory page')]}
    },
    'operator': {
        'th': {'title': 'คู่มือผู้ใช้งาน WMS', 'role': 'Operator (พนักงานเบิกสินค้า)', 'tag': 'ค้นหา สร้างใบเบิก และติดตามผล',
               'summary': 'ใช้สำหรับพนักงานที่ต้องเบิกอะไหล่หรือสินค้าไปใช้งานโครงการ โดยต้องรอ Manager หรือ Admin อนุมัติ',
               'allowed': ['สิทธิ์ทั้งหมดของ Viewer', 'เพิ่มสินค้าในตะกร้าใบเบิก', 'ระบุโครงการและส่งใบเบิกหลายรายการพร้อมกัน', 'ยกเลิกเฉพาะใบเบิกของตนที่ยังรออนุมัติ', 'ติดตามผลและหมายเหตุจากผู้อนุมัติ'],
               'steps': [('เลือกสินค้า', ['เปิด “สินค้าคงคลัง”', 'ค้นหา/สแกนสินค้าและตรวจสอบยอดคงเหลือ', 'กด “เพิ่ม” ในแถวสินค้าที่ต้องการ']), ('จัดทำใบเบิก', ['กด “ตะกร้าใบเบิก” มุมล่างขวา', 'ปรับจำนวนด้วยปุ่ม − / + หรือกรอกจำนวน', 'เลือกโครงการ แล้วกด “ส่งใบเบิก”']), ('ติดตามผล', ['เปิด Dashboard หรือกระดิ่งแจ้งเตือน', 'อ่านสถานะ Approved, Partial หรือ Rejected พร้อมหมายเหตุ', 'เมื่อขึ้น “มารับสินค้าได้” ให้รับสินค้าจากคลัง'])],
               'notes': ['ขอเบิกเกินยอด “เบิกได้” ของโครงการที่เลือกไม่ได้ (ไม่ใช่ยอดคงเหลือรวม)', 'สต็อกจะลดเมื่อผู้ดูแลกด “รับแล้ว” เท่านั้น การอนุมัติเป็นเพียงการกันของไว้ให้', 'ถ้าไม่มารับของนาน ผู้ดูแลอาจยกเลิกการจองและคืนของเข้าสต็อก โดยจะแจ้งเหตุผลให้ทราบ', 'ยกเลิกได้เฉพาะสถานะ Pending และเฉพาะใบของตนเอง'],
               'shots': [('operator-inventory.png', 'ภาพที่ 1: เลือกสินค้าและเพิ่มลงใบเบิก'), ('operator-request-cart.png', 'ภาพที่ 2: ตรวจสอบจำนวน เลือกโครงการ และส่งใบเบิก')]},
        'en': {'title': 'WMS User Manual', 'role': 'Operator', 'tag': 'Find items, submit requests, and track results',
               'summary': 'For staff who withdraw parts or supplies for a project. Every request must be reviewed by a Manager or Admin.',
               'allowed': ['All Viewer capabilities', 'Add items to a withdrawal cart', 'Choose a project and submit multiple items in one request', 'Cancel only your own pending request', 'Track results and approver notes'],
               'steps': [('Select items', ['Open “Inventory”', 'Search/scan an item and check available balance', 'Select “Add” on the required item row']), ('Create a request', ['Select “Request cart” at the lower-right', 'Adjust quantities with − / + or enter a quantity', 'Choose a project and select “Submit request”']), ('Track the result', ['Open Dashboard or the notification bell', 'Read Approved, Partial, or Rejected status and the note', 'When “Ready for pickup” appears, collect the item from the warehouse'])],
               'notes': ['A request cannot exceed the Available figure for the selected project (not the total balance).', 'Stock decreases only when an administrator selects “Picked up”; approval merely holds the goods for you.', 'If goods are not collected for a long time, an administrator may cancel the hold and return them to stock, with a stated reason.', 'Only your own Pending request can be cancelled.'],
               'shots': [('operator-inventory.png', 'Figure 1: Select an item and add it to the request'), ('operator-request-cart.png', 'Figure 2: Confirm quantity, choose project, and submit')]}
    },
    'manager': {
        'th': {'title': 'คู่มือผู้ใช้งาน WMS', 'role': 'Manager (ผู้จัดการคลัง)', 'tag': 'ควบคุมข้อมูลสินค้า สต็อก และใบเบิก',
               'summary': 'ใช้สำหรับผู้ดูแลคลังที่จัดการสินค้า รับเข้า ปรับยอด อนุมัติใบเบิก และวิเคราะห์ข้อมูลได้ โดยไม่มีสิทธิ์จัดการบัญชีผู้ใช้',
               'allowed': ['สิทธิ์ทั้งหมดของ Operator', 'เพิ่ม แก้ไข รับเข้า ปรับยอด และปิดใช้งานสินค้า', 'นำเข้า/ส่งออกข้อมูลสินค้า CSV', 'จัดการหมวดหมู่และโครงการ', 'อนุมัติ ปฏิเสธ และยืนยันส่งมอบใบเบิก', 'ดู Analytics และ AI forecast'],
               'steps': [('จัดการสินค้า', ['เปิด “รายการอะไหล่”', 'กดเพิ่มสินค้า หรือเลือกแก้ไข/รับเข้า/ปรับยอดบนการ์ดสินค้า', 'ตรวจสอบ SKU หมวดหมู่ หน่วย ผู้ขาย ราคา สต็อกขั้นต่ำ และบันทึก']), ('พิจารณาใบเบิกและส่งมอบ', ['เปิด Dashboard และกด “ตรวจสอบ” ที่ใบ Pending', 'ตั้งจำนวนอนุมัติรายรายการ — อนุมัติได้ไม่เกินยอดที่เบิกได้ของโครงการนั้น (ถ้ามีพื้นที่จัดเตรียม จะจำกัดตามจำนวนที่จัดเตรียมไว้)', 'ถ้าโควตาไม่พอ ให้เติมของเข้าพื้นที่จัดเตรียมของโครงการก่อน แล้วค่อยอนุมัติ', 'กรอกเหตุผลหากไม่อนุมัติครบหรือปฏิเสธ แล้วบันทึกผล — ตอนนี้ระบบจะกันของไว้ให้ ยังไม่ตัดสต็อก', 'เมื่อผู้ขอมารับของ กด “รับแล้ว” ระบบจึงตัดสต็อกจริง โดยหยิบจากพื้นที่จัดเตรียมของโครงการนั้นก่อน', 'ถ้าไม่มีผู้มารับนานเกินไป กด “ยกเลิกจอง” พร้อมระบุเหตุผล เพื่อคืนของเข้าสต็อก']), ('วิเคราะห์และส่งออก', ['เปิด “วิเคราะห์คลัง” เพื่อดูแนวโน้มและรายการที่เบิกมาก', 'ใช้ CSV import/export สำหรับอัปเดตข้อมูลจำนวนมาก', 'ติดตามสต็อกต่ำและปรับยอดจากผลนับจริง'])],
               'notes': ['อนุมัติบางส่วนหรืออนุมัติศูนย์ต้องมีเหตุผล', 'การอนุมัติเป็นการ “กันของไว้” ยังไม่ตัดสต็อก — สต็อกถูกตัดจริงตอนกด “รับแล้ว”', 'ยกเลิกการจองต้องระบุเหตุผลทุกครั้ง ระบบจะแจ้งผู้ขอเบิกพร้อมเหตุผลนั้น', 'การปรับยอดตามการนับจริงยังตัดสต็อกทันทีเหมือนเดิม ไม่ผ่านระบบจอง', 'ควรใช้ “ปิดใช้งาน” ก่อนพิจารณาลบสินค้าถาวร'],
               'shots': [('manager-dashboard.png', 'ภาพที่ 1: Dashboard สำหรับตรวจสอบใบเบิก'), ('manager-products.png', 'ภาพที่ 2: รายการอะไหล่และคำสั่งจัดการสินค้า'), ('admin-analysis.png', 'ภาพที่ 3: หน้าวิเคราะห์คลัง (สิทธิ์ Manager และ Admin)')]},
        'en': {'title': 'WMS User Manual', 'role': 'Manager', 'tag': 'Control catalog, stock, and withdrawal requests',
               'summary': 'For warehouse supervisors who manage products, inbound stock, stock adjustments, approvals, and analysis. Managers cannot manage user accounts.',
               'allowed': ['All Operator capabilities', 'Create, edit, receive, adjust, and deactivate products', 'Import/export product data as CSV', 'Manage categories and projects', 'Approve, reject, and confirm pickup', 'View Analytics and AI forecast'],
               'steps': [('Manage products', ['Open “Products”', 'Add a product or use edit/receive/adjust actions on a product card', 'Check SKU, category, unit, vendor, price, minimum stock, then save']), ('Review requests and hand over', ['Open Dashboard and select “Review” for a Pending request', 'Set an approved quantity per item — it cannot exceed the project’s available figure (capped by the staging area when one exists)', 'If the quota is short, add stock to that project’s staging area first, then approve', 'Enter a reason for incomplete/rejected quantities and save — the goods are now held, stock is not yet deducted', 'When the requester collects, select “Picked up”; only then is stock deducted, taken from that project’s staging area first', 'If nobody collects, select “Cancel hold” with a reason to return the goods to stock']), ('Analyze and export', ['Open “Analytics” for trends and high-demand items', 'Use CSV import/export for bulk updates', 'Monitor low stock and adjust stock from a physical count'])],
               'notes': ['Partial or zero approval requires a reason.', 'Approval only holds the goods; stock is deducted when you select “Picked up”.', 'Cancelling a hold always requires a reason, which is sent to the requester.', 'A physical-count adjustment still deducts stock immediately and does not go through the hold system.', 'Use Deactivate before considering permanent deletion.'],
               'shots': [('manager-dashboard.png', 'Figure 1: Manager dashboard for reviewing requests'), ('manager-products.png', 'Figure 2: Product catalog management actions'), ('admin-analysis.png', 'Figure 3: Analytics page (Manager and Admin access)')]}
    },
    'admin': {
        'th': {'title': 'คู่มือผู้ใช้งาน WMS', 'role': 'Admin (ผู้ดูแลระบบ)', 'tag': 'ดูแลผู้ใช้ พร้อมสิทธิ์จัดการคลังครบถ้วน',
               'summary': 'ใช้สำหรับผู้ดูแลระบบ มีสิทธิ์ทั้งหมดของ Manager และดูแลสถานะ/บทบาทของผู้ใช้งาน',
               'allowed': ['สิทธิ์ทั้งหมดของ Manager', 'ดูรายชื่อและกรองสถานะผู้ใช้', 'อนุมัติ ปฏิเสธ ระงับ หรือคืนสิทธิ์บัญชี', 'กำหนด Role: Admin, Manager, Operator, Viewer', 'ลบผู้ใช้ออกจากระบบ (ยกเว้นตนเอง)', 'รับแจ้งเตือนผู้สมัครใหม่'],
               'steps': [('อนุมัติผู้สมัคร', ['เปิด “จัดการผู้ใช้งาน”', 'กรองสถานะ Pending หากต้องการ', 'ตรวจสอบชื่อผู้ใช้และอีเมล แล้วกด “อนุมัติ” หรือ “ปฏิเสธ”']), ('กำหนดสิทธิ์', ['เลือก Role จากรายการของผู้ใช้', 'กำหนด Operator สำหรับผู้ขอเบิก หรือ Manager สำหรับผู้ดูแลคลัง', 'แจ้งผู้ใช้ให้เข้าสู่ระบบใหม่หากเพิ่งเปลี่ยนสิทธิ์']), ('ดูแลคลัง', ['ใช้ Dashboard ตรวจสอบใบเบิกและสต็อกต่ำ', 'ใช้ “รายการอะไหล่” จัดการสินค้าและ CSV', 'ใช้ “วิเคราะห์คลัง” ตรวจสอบแนวโน้มและแผนสั่งซื้อ']), ('พิจารณาใบเบิกและส่งมอบ (เหมือนผู้จัดการ)', ['กด “ตรวจสอบ” ที่ใบ Pending แล้วตั้งจำนวนอนุมัติรายรายการ', 'อนุมัติได้ไม่เกินยอดที่เบิกได้ของโครงการนั้น — ถ้าโครงการมีพื้นที่จัดเตรียม จะถูกจำกัดตามโควตาที่จัดเตรียมไว้ ต้องเติมของเข้าพื้นที่ก่อนจึงอนุมัติเพิ่มได้', 'การอนุมัติเป็นการกันของไว้ให้เท่านั้น ยังไม่ตัดสต็อก', 'กด “รับแล้ว” เมื่อผู้ขอมารับของ ระบบจึงตัดสต็อกจริง โดยหยิบจากพื้นที่จัดเตรียมของโครงการนั้นก่อน', 'ถ้าไม่มีผู้มารับนานเกินไป กด “ยกเลิกจอง” พร้อมระบุเหตุผล เพื่อคืนของเข้าสต็อก'])],
               'notes': ['ลบบัญชีตนเองไม่ได้', 'ระบบไม่อนุญาตให้ลดสิทธิ์/ระงับ/ลบ Admin ที่ใช้งานอยู่คนสุดท้าย', 'สต็อกถูกตัดตอนกด “รับแล้ว” ไม่ใช่ตอนอนุมัติ — การอนุมัติเป็นการกันของไว้ให้เท่านั้น', 'การเปลี่ยนสถานะบัญชีจะส่งอีเมลแจ้งผู้สมัครเมื่อระบบอีเมลถูกตั้งค่า'],
               'shots': [('admin-dashboard.png', 'ภาพที่ 1: Dashboard ของ Admin'), ('admin-users.png', 'ภาพที่ 2: หน้าจัดการผู้ใช้งาน'), ('admin-products.png', 'ภาพที่ 3: หน้าจัดการสินค้า'), ('admin-analysis.png', 'ภาพที่ 4: หน้าจอวิเคราะห์คลัง')]},
        'en': {'title': 'WMS User Manual', 'role': 'Admin', 'tag': 'Administer users with full warehouse access',
               'summary': 'For system administrators. Admins have all Manager permissions and control user status and roles.',
               'allowed': ['All Manager capabilities', 'View users and filter by status', 'Approve, reject, suspend, or reactivate accounts', 'Assign Admin, Manager, Operator, or Viewer roles', 'Delete users other than yourself', 'Receive new-registration notifications'],
               'steps': [('Approve applicants', ['Open “User management”', 'Filter to Pending when needed', 'Review username and email, then select “Approve” or “Reject”']), ('Assign access', ['Select a role from the user role list', 'Assign Operator to requesters or Manager to warehouse supervisors', 'Ask the user to sign in again after a new role is assigned']), ('Administer the warehouse', ['Use Dashboard to review requests and low-stock items', 'Use “Products” for catalog and CSV administration', 'Use “Analytics” for trend and replenishment review']), ('Review requests and hand over (same as Manager)', ['Select “Review” on a Pending request and set an approved quantity per item', 'Approval cannot exceed the project’s available figure — where the project has a staging area, the quota caps it, so top the area up first to approve more', 'Approval only holds the goods; stock is not yet deducted', 'Select “Picked up” when the requester collects; only then is stock deducted, taken from that project’s staging area first', 'If nobody collects, select “Cancel hold” with a reason to return the goods to stock'])],
               'notes': ['You cannot delete your own account.', 'The system prevents removing, suspending, or demoting the last active Admin.', 'Stock is deducted on “Picked up”, not on approval — approving only holds the goods.', 'A status change emails the applicant when email is configured.'],
               'shots': [('admin-dashboard.png', 'Figure 1: Admin dashboard'), ('admin-users.png', 'Figure 2: User management'), ('admin-products.png', 'Figure 3: Product management'), ('admin-analysis.png', 'Figure 4: Warehouse analytics')]}
    }
}

# ผังคลังสินค้า (Storage Map) — เนื้อหาแยก 2 ระดับ แล้วผสมเข้าทุก role ด้านล่าง
# view = ดูอย่างเดียว (Viewer, Operator) | edit = แก้ไขผังได้ (Manager, Admin)
STORAGE = {
    'view': {
        'th': {
            'allowed': ['เปิดเมนู “ผังคลัง” เพื่อดูผังห้องเก็บของและชั้นวางทั้งหมด',
                        'ค้นหาตำแหน่งจัดเก็บด้วยรหัสสินค้า (SKU) ชื่อห้อง หรือชื่อชั้นวาง',
                        'เปิดผังชั้นวางเพื่อดูรายการสินค้าแยกตามเลเวล'],
            'steps': [('อ่านยอดคงเหลือ กันไว้ และเบิกได้', [
                'เปิดเมนู “สินค้าคงคลัง” แล้วเลือกโครงการที่จะนำสินค้าไปใช้จากช่องเลือกโครงการก่อนเสมอ',
                'ถ้ายังไม่เลือกโครงการ ระบบจะขึ้นแถบเตือน และแสดงเฉพาะ “ของกลาง” ที่ยังไม่ได้กันไว้ให้โครงการใด',
                'อ่านตัวเลข 3 ช่อง — คงเหลือ คือของที่มีอยู่จริงทั้งหมด, กันไว้ คือของที่ถูกจองให้ใบเบิกที่อนุมัติแล้วรอรับ, เบิกได้ คือจำนวนที่ขอเบิกได้จริงตอนนี้',
                'ถ้าช่องเบิกได้มีป้าย “โควตา” แปลว่าโครงการนี้มีพื้นที่จัดเตรียมของตัวเอง เบิกได้ไม่เกินจำนวนที่จัดเตรียมไว้',
                'ถ้าต้องการเบิกเกินโควตา ให้แจ้งผู้จัดการเติมของเข้าพื้นที่จัดเตรียมก่อน']),
                ('ค้นหาตำแหน่งจัดเก็บสินค้า', [
                'เลือกเมนู “ผังคลัง” ระบบจะแสดงผังทั้งคลังที่ระดับซูม 80%',
                'ห้องเก็บของแสดงเป็นกรอบสีน้ำเงิน ส่วนชั้นวางแสดงเป็นบล็อกสีฟ้า',
                'พิมพ์รหัสสินค้า (SKU) ชื่อห้อง หรือชื่อชั้นวางในช่องค้นหา แล้วกด “ค้นหา”',
                'ระบบจะเลื่อนผังไปยังตำแหน่งนั้น เปิดผังชั้นวาง และไฮไลต์เลเวลที่จัดเก็บสินค้า',
                'คลิกที่แถบเลเวลเพื่อดูรายการสินค้าบนเลเวลนั้น พร้อมรูปภาพ รหัสสินค้า ชื่อสินค้า หมวดหมู่ และยอดคงเหลือ',
                'ปรับมุมมองด้วยล้อเมาส์ (หมุนขึ้น-ลง = ซูมเข้า-ออก) หรือบนมือถือใช้สองนิ้วหุบ-กาง ซูมได้ 20%-400%',
                'เลื่อนผังบนมือถือด้วยการลากนิ้วเดียว บนคอมพิวเตอร์กดค้าง Space แล้วลาก',
                'ถ้าหลงทางให้กดปุ่ม “พอดีหน้าจอ” ในแถบล่างเพื่อกลับมาเห็นผังทั้งหมด'])],
            'notes': ['ผังคลังของบัญชีนี้เป็นแบบดูอย่างเดียว การเพิ่ม ย้าย หรือแก้ไขผังเป็นสิทธิ์ของ Manager และ Admin',
                      'สินค้าที่ยังไม่ได้กำหนดตำแหน่งจัดเก็บจะค้นหาในผังคลังไม่พบ',
                      'ก่อนดูยอด “เบิกได้” ให้เลือกโครงการที่หน้าสินค้าคงคลังก่อนเสมอ ยอดที่เบิกได้ของแต่ละโครงการไม่เท่ากัน',
                      'สต็อกจะถูกตัดจริงตอนผู้ดูแลกด “รับแล้ว” ไม่ใช่ตอนอนุมัติ ระหว่างรออนุมัติถึงรับของ ระบบจะกันของไว้ให้']
        },
        'en': {
            'allowed': ['Open “Storage map” to view all storage rooms and racks',
                        'Search for a storage location by SKU, room name, or rack name',
                        'Open a rack blueprint to see items grouped by level'],
            'steps': [('Read balance, held, and available', [
                'Open “Inventory” and choose the project you are withdrawing for before anything else',
                'Without a project selected, a banner appears and only unallocated “free” stock is shown',
                'Read the three figures — Balance is everything physically in stock, Held is reserved for approved requests awaiting pickup, Available is what you can request right now',
                'A “quota” badge on Available means this project has its own staging area and cannot exceed the prepared amount',
                'To request more than the quota, ask a manager to add stock to the staging area first']),
                ('Find where an item is stored', [
                'Open “Storage map”; the full layout opens at 80% zoom',
                'Storage rooms appear as blue outlines and racks as light-blue blocks',
                'Enter a SKU, room name, or rack name in the search box and select “Search”',
                'The map pans to that location, opens the rack blueprint, and highlights the storage level',
                'Select a level row to list every item on it with image, SKU, name, category, and balance',
                'Zoom with the mouse wheel (scroll up/down) or a two-finger pinch on mobile; the range is 20%-400%',
                'Pan by dragging with one finger on mobile, or by holding Space and dragging on a computer',
                'Press “Fit to screen” in the bottom toolbar to see the whole layout again'])],
            'notes': ['The storage map is read-only for this role; adding, moving, or editing the layout requires Manager or Admin.',
                      'Items without an assigned storage location cannot be found on the map.',
                      'Select a project on the Inventory page before reading the “Available” figure — each project has its own quota.',
                      'Stock is deducted when an administrator selects “Picked up”, not on approval; goods are held for you in between.']
        }
    },
    'edit': {
        'th': {
            'allowed': ['เปิดเมนู “ผังคลัง” ค้นหาตำแหน่งจัดเก็บด้วยรหัสสินค้า (SKU) และดูสินค้าแยกตามเลเวลของชั้นวาง',
                        'สร้างและแก้ไขผังคลัง: เพิ่มห้อง ชั้นวาง และสัญลักษณ์ (กำแพง ประตู เส้นแบ่ง ข้อความ)',
                        'กำหนดขนาดห้องและชั้นวางตั้งแต่ตอนสร้าง และปรับขนาดภายหลังได้ทุกเมื่อ',
                        'ลาก หมุนอิสระ ล็อก จัดกลุ่ม และจัดลำดับการซ้อน (Layer) ขององค์ประกอบบนผัง',
                        'เผยแพร่ผัง (Publish) พร้อมเก็บประวัติเวอร์ชัน และกู้คืนรายการที่ลบจากถังขยะ',
                        'ผูกสินค้าเข้าชั้นวาง/เลเวล และดูรายการสินค้าที่ยังไม่ระบุตำแหน่ง',
                        'ดูเส้นทางเดินหยิบของทั้งใบเบิกบนผัง และข้ามไปยังตำแหน่งจริงของสินค้าแต่ละรายการ',
                        'สร้างพื้นที่จัดเตรียมผูกกับโครงการ และกำหนดรายการ/จำนวนสินค้าที่กันไว้ให้โครงการนั้น',
                        'ยกเลิกการจองของใบที่อนุมัติแล้วแต่ไม่มีผู้มารับ เพื่อคืนของเข้าสต็อก'],
            'steps': [('จัดการผังคลัง', [
                'เลือกเมนู “ผังคลัง” ระบบจะแสดงผังทั้งคลังที่ระดับซูม 80% แล้วสลับเป็น “โหมดแก้ไข”',
                'ห้องเก็บของแสดงเป็นกรอบสีน้ำเงิน ชั้นวางแสดงเป็นบล็อกสีฟ้า และเปลี่ยนเป็นสีเหลืองเมื่อความจุเกิน 70% หรือสีแดงเมื่อเกิน 90%',
                'เลือกคลังที่ต้องการจากรายการด้านบน หรือกดปุ่มจัดการคลังเพื่อเพิ่ม/แก้ไขคลัง',
                'เลือกเครื่องมือเพิ่มห้องหรือชั้นวาง แล้ว “ลากบนผังเพื่อกำหนดขนาดที่ต้องการ” ระบบจะแสดงขนาดกว้าง × สูงระหว่างลาก (คลิกเฉยๆ = ใช้ขนาดมาตรฐาน)',
                'ตั้งชื่อ จำนวนเลเวล และความจุของชั้นวางในแผงคุณสมบัติด้านข้าง',
                'ปรับขนาดภายหลังได้ 2 ทาง: ลากจุดวงกลมที่มุมขวาล่างขององค์ประกอบ หรือกรอกตัวเลขในช่อง “กว้าง/สูง” ของแผงคุณสมบัติ',
                'เมื่อจัดผังเรียบร้อย กด “Publish” เพื่อบันทึกเป็นเวอร์ชันใช้งานจริง']),
                ('หมุนองค์ประกอบและวาดกำแพง', [
                'เลือกองค์ประกอบ จะเห็นป้ายบอกองศาปัจจุบันอยู่เหนือองค์ประกอบนั้น',
                'ลากปุ่มหมุนวงกลมมุมขวาบนเพื่อหมุนอิสระทีละ 1 องศา โดยองศาจะอัปเดตให้เห็นตลอดที่ลาก',
                'กดปุ่ม Shift ค้างไว้ระหว่างลาก เพื่อล็อกการหมุนทีละ 15 องศา (ได้มุมฉากหรือ 45 องศาแบบตรงพอดี)',
                'ถ้าต้องการองศาที่แน่นอน ให้กรอกตัวเลขลงช่อง “หมุน (องศา)” ในแผงคุณสมบัติได้โดยตรง',
                'กำแพงและเส้นแบ่งจะหมุนโดยยึดปลายด้านที่เริ่มลากไว้กับที่ ปลายอีกด้านเท่านั้นที่กวาดไป ทำให้ต่อกำแพงเป็นมุมห้องได้ง่าย']),
                ('กำหนดตำแหน่งจัดเก็บสินค้า', [
                'กดปุ่ม “ยังไม่ระบุตำแหน่ง” บนแถบด้านบน เพื่อดูรายการสินค้าที่ยังไม่มีตำแหน่งจัดเก็บ',
                'ค้นหาสินค้า เลือกชั้นวางและเลเวล แล้วกดบันทึก',
                'หรือเปิดผังชั้นวาง เข้าไปที่เลเวลที่ต้องการ แล้วกด “เพิ่มสินค้าเข้าเลเวลนี้”',
                'กดปุ่มถังขยะท้ายรายการสินค้า เพื่อถอนสินค้าออกจากตำแหน่งเดิม']),
                ('จัดพื้นที่จัดเตรียมของโครงการ', [
                'พื้นที่จัดเตรียมคือโซนบนผังที่วางของกับพื้นได้โดยตรง ไม่ต้องมีชั้นวาง ใช้กันของไว้ให้โครงการใดโครงการหนึ่ง',
                'เลือกเครื่องมือ “ห้อง” ติ๊ก “พื้นที่จัดเตรียม” แล้วเลือกโครงการจากรายการ จากนั้นลากกำหนดขนาดบนผัง',
                'พื้นที่จัดเตรียมจะแสดงเป็นกรอบสีเหลืองอำพัน พร้อมชื่อโครงการและจำนวนของข้างใน',
                'ย้ายของเข้าพื้นที่: เปิดผังชั้นวาง เข้าเลเวลที่เก็บของ กดปุ่มกล่องท้ายแถวสินค้า เลือกพื้นที่และจำนวน แล้วกด “ย้าย”',
                'แก้รายการ/จำนวนในพื้นที่: ดับเบิลคลิกที่พื้นที่บนผัง แล้วแก้ตัวเลขหรือกดถังขยะเพื่อเอาสินค้าออก',
                'ของที่อยู่ในพื้นที่จัดเตรียมยังนับเป็นสต็อกอยู่ แต่ถูกกันไว้ให้โครงการนั้น โครงการอื่นเบิกไม่ได้']),
                ('ตรวจตำแหน่งจัดเก็บก่อนเตรียมสินค้าตามใบเบิก', [
                'เปิด Dashboard แล้วคลิกใบเบิกในตาราง “คำขอเบิกรอดำเนินการ / รอส่งมอบ”',
                'ดูคอลัมน์ “ตำแหน่งจัดเก็บ” ของแต่ละรายการ ซึ่งแสดงชื่อชั้นวางและเลเวล',
                'กดปุ่ม “ดูทั้งใบบนผังคลัง” เพื่อดูเส้นทางเดินหยิบของทั้งใบ ระบบจะเรียงลำดับจุดแวะให้เดินน้อยที่สุด',
                'ชั้นวางที่ต้องแวะจะมีป้ายตัวเลขลำดับกำกับบนผัง ติ๊กรายการที่หยิบแล้วในแผงด้านขวา ป้ายจะเปลี่ยนเป็นเครื่องหมายถูก',
                'เตรียมสินค้าตามตำแหน่งจริง แล้วกลับมากด “รับแล้ว” เมื่อผู้ขอมารับของ'])],
            'notes': ['ลบชั้นวางที่ยังมีสินค้าอยู่ไม่ได้ ต้องย้ายสินค้าออกก่อน และลดจำนวนเลเวลที่มีสินค้าอยู่ไม่ได้เช่นกัน',
                      'ต้องปลดล็อกองค์ประกอบที่ล็อกไว้ก่อน จึงจะย้าย ปรับขนาด หรือหมุนได้',
                      'ขนาดขั้นต่ำของห้องคือ 80 × 60 และของชั้นวางคือ 60 × 40',
                      'ถ้าหมุนกำแพงแล้วปลายจะพ้นขอบผัง ระบบจะดันกลับเข้าผังให้ ตำแหน่งจึงอาจขยับเล็กน้อย',
                      'สินค้าที่ยังไม่ระบุตำแหน่งจะไม่ปรากฏบนผังและค้นหาด้วย SKU ไม่พบ',
                      'ลบพื้นที่จัดเตรียมที่ยังมีสินค้าอยู่ไม่ได้ และเปลี่ยนกลับเป็นห้องธรรมดาไม่ได้จนกว่าจะย้ายของออกหมด',
                      'ห้องที่มีชั้นวางอยู่แล้ว เปลี่ยนเป็นพื้นที่จัดเตรียมไม่ได้ ต้องย้ายชั้นวางออกก่อน',
                      'รายการที่ลบจะถูกย้ายไปถังขยะและกู้คืนได้จากเมนู “Version / Trash”']
        },
        'en': {
            'allowed': ['Open “Storage map”, search a location by SKU, and view items by rack level',
                        'Build and edit the layout: add rooms, racks, and markers (walls, doors, dividers, text)',
                        'Set room and rack size while drawing, and resize them again at any time',
                        'Move, freely rotate, lock, group, and reorder layers of layout components',
                        'Publish the layout with version history and restore deleted items from the trash',
                        'Assign items to a rack level and review items that still have no location',
                        'View the pick route for a whole request and jump to each item’s actual location',
                        'Create project-linked staging areas and set which items and quantities are held for each project',
                        'Cancel a hold on an approved request that was never collected, returning the goods to stock'],
            'steps': [('Manage the storage layout', [
                'Open “Storage map”; the layout opens at 80% zoom. Switch to “Edit mode”',
                'Storage rooms appear as blue outlines and racks as light-blue blocks, turning amber above 70% capacity and red above 90%',
                'Select a warehouse from the list, or use the manage button to add or edit one',
                'Pick the room or rack tool, then drag on the canvas to set the size you want — the width × height is shown while dragging (a plain click uses the default size)',
                'Set the rack name, level count, and capacity in the side properties panel',
                'Resize later in two ways: drag the round handle at the lower-right corner, or type exact numbers into the Width/Height fields',
                'When the layout is ready, select “Publish” to save it as the live version']),
                ('Rotate components and draw walls', [
                'Select a component to see its current angle on a badge above it',
                'Drag the round rotate handle at the upper-right to rotate freely, one degree at a time, with the angle updating as you drag',
                'Hold Shift while dragging to snap the rotation to 15° steps for exact right angles or 45° diagonals',
                'For an exact angle, type the value into the “Rotation (degrees)” field in the properties panel',
                'Walls and dividers pivot around the end you started drawing from, so only the far end sweeps — this makes joining walls into corners much easier']),
                ('Assign storage locations', [
                'Select “Unassigned” in the top bar to list items that have no storage location yet',
                'Search for an item, choose a rack and level, then save',
                'Or open a rack blueprint, go to the level you want, and select “Add item to this level”',
                'Use the trash button at the end of an item row to remove it from its current location']),
                ('Set up a project staging area', [
                'A staging area is a zone on the map where goods sit directly on the floor, with no rack, reserved for one project',
                'Pick the room tool, tick “Staging area”, choose the project, then drag to size it on the map',
                'Staging areas appear as amber outlines showing the project name and the quantity inside',
                'To move goods in: open a rack blueprint, go to the level, select the box button on the item row, choose the area and quantity, then Move',
                'To edit contents: double-click the area on the map, then change the number or use the bin button to remove an item',
                'Goods in a staging area still count as stock but are held for that project only — no other project can withdraw them']),
                ('Check locations before preparing a request', [
                'Open Dashboard and select a request in the pending/awaiting-pickup table',
                'Review the “Storage location” column showing rack name and level for each item',
                'Select “View whole request on map” to see the pick route, ordered to minimise walking',
                'Each rack to visit is numbered on the map; tick items off in the right-hand panel and the badge turns into a check mark',
                'Pick the items from their actual location, then select “Picked up” when the requester collects them'])],
            'notes': ['A rack holding items cannot be deleted, and its level count cannot be reduced below occupied levels.',
                      'Locked components must be unlocked before they can be moved, resized, or rotated.',
                      'Minimum size is 80 × 60 for a room and 60 × 40 for a rack.',
                      'If a rotation would push a wall past the canvas edge, it is pulled back inside, so its position may shift slightly.',
                      'Items without a storage location do not appear on the map and cannot be found by SKU search.',
                      'A staging area holding goods cannot be deleted or converted back to a normal room until it is emptied.',
                      'A room that already contains racks cannot be turned into a staging area; move the racks out first.',
                      'Deleted components move to the trash and can be restored from “Version / Trash”.']
        }
    }
}

# ผสมเนื้อหาผังคลังเข้าไปในทุก role (Manager/Admin ได้ระดับแก้ไข ที่เหลือได้ระดับดูอย่างเดียว)
for _role, _langs in ROLE.items():
    _tier = 'edit' if _role in ('manager', 'admin') else 'view'
    for _lang, _data in _langs.items():
        _extra = STORAGE[_tier][_lang]
        _data['allowed'] = _data['allowed'] + _extra['allowed']
        _data['steps'] = _data['steps'] + _extra['steps']
        _data['notes'] = _data['notes'] + _extra['notes']


def set_font(run, size=None, bold=None, color=None):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = color

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    set_font(run, 9, color=MUTED)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

def init_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles['Normal']; normal.font.name = 'Arial'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18
    for n, sz, col, before, after in [('Heading 1',16,ACCENT,16,8),('Heading 2',13,ACCENT,12,6),('Heading 3',11.5,DARK,8,4)]:
        s=styles[n]; s.font.name='Arial'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial'); s.font.size=Pt(sz); s.font.color.rgb=col; s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    hp=sec.header.paragraphs[0]; hp.text='WMS iCreativeSystems | User Manual'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs: set_font(r,8.5,color=MUTED)
    add_page_number(sec.footer.paragraphs[0])
    return doc

def para(doc, text='', style=None, bold_lead=None):
    p=doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_font(r,bold=True,color=DARK)
        r=p.add_run(text[len(bold_lead):]); set_font(r)
    else:
        r=p.add_run(text); set_font(r)
    return p

def bullet(doc, text):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); set_font(r)

def step(doc, text):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(text); set_font(r)

def image(doc, filename, caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(SHOTS/filename), width=Inches(6.55))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(8)
    r=c.add_run(caption); set_font(r,9,color=MUTED)

def build(role, lang, data):
    doc=init_doc()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(72); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(data['title']); set_font(r,25,bold=True,color=ACCENT)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8)
    r=p.add_run(data['role']); set_font(r,17,bold=True,color=DARK)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(34)
    r=p.add_run(data['tag']); set_font(r,11,color=MUTED)
    para(doc, 'WMS iCreativeSystems | Version 2.0 | 26 August 2026')
    doc.add_page_break()
    para(doc, '1. Overview' if lang=='en' else '1. ภาพรวม', style='Heading 1')
    para(doc, data['summary'])
    para(doc, 'What you can do' if lang=='en' else 'สิ่งที่สามารถทำได้', style='Heading 2')
    for x in data['allowed']: bullet(doc,x)
    para(doc, 'Before you start' if lang=='en' else 'ก่อนเริ่มใช้งาน', style='Heading 2')
    for x in (['Use your own account; one account can be active on one device at a time.', 'Use the profile menu for password, email, profile image, theme, and notification settings.'] if lang=='en' else ['ใช้บัญชีของตนเอง โดย 1 บัญชีใช้งานได้ทีละ 1 อุปกรณ์', 'ใช้เมนูโปรไฟล์สำหรับเปลี่ยนรหัสผ่าน อีเมล รูปโปรไฟล์ ธีม และการแจ้งเตือน']): bullet(doc,x)
    for title, items in data['steps']:
        para(doc, title, style='Heading 1')
        for x in items: step(doc,x)
    para(doc, 'Screens from the live test system' if lang=='en' else 'ภาพประกอบจากระบบทดสอบจริง', style='Heading 1')
    for filename, caption in data['shots']:
        image(doc, filename, caption)
    para(doc, 'Important notes' if lang=='en' else 'ข้อควรระวัง', style='Heading 1')
    for x in data['notes']: bullet(doc,x)
    para(doc, 'Support' if lang=='en' else 'การขอความช่วยเหลือ', style='Heading 1')
    para(doc, 'Contact your WMS Administrator if access, stock data, or a workflow result appears incorrect.' if lang=='en' else 'ติดต่อผู้ดูแลระบบ WMS หากสิทธิ์ ข้อมูลสต็อก หรือผลของกระบวนการทำงานไม่ถูกต้อง')
    name=f'WMS_User_Manual_{role.capitalize()}_{"EN" if lang=="en" else "TH"}.docx'
    doc.save(OUT/name)

if __name__ == '__main__':
    for role, langs in ROLE.items():
        for lang, data in langs.items(): build(role,lang,data)
    print('Created', len(list(OUT.glob('*.docx'))), 'manuals in', OUT)
