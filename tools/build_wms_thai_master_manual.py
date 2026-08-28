"""Build the comprehensive Thai WMS user guideline as an editable DOCX.

Design system: compact_reference_guide + editorial_cover, adapted for Thai
readability with Arial as the named font override.
"""
from __future__ import annotations

import argparse
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "user-guidelines"
SHOT_DIR = ROOT / "output" / "manual-screenshots"
LOGO = ROOT / "public" / "icons" / "ICS.png"
OUT_FILE = OUT_DIR / "WMS_User_Guideline_TH_All_Roles_Final.docx"

FONT = "Arial"  # named override for reliable Thai rendering in Word/LibreOffice
PAGE_W_DXA = 12240  # US Letter 8.5 in
PAGE_H_DXA = 15840  # US Letter 11 in
CONTENT_W_DXA = 9360  # 6.5 in after 1 in left/right margins
TABLE_INDENT_DXA = 120

BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
TEAL = "0D8ABC"
INK = "243447"
MUTED = "5F6B7A"
PALE_BLUE = "EAF3FA"
PALE_TEAL = "EAF8FA"
PALE_YELLOW = "FFF4D6"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF7EE"
TABLE_HEADER = "E8EEF5"
WHITE = "FFFFFF"
LIGHT_LINE = "D8E2EB"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size: float = 11, bold: bool = False, color: str = INK,
                 italic: bool = False) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), FONT)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "th-TH")
    lang.set(qn("w:eastAsia"), "th-TH")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in edges.items():
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 6)))
        el.set(qn("w:color"), spec.get("color", LIGHT_LINE))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: Sequence[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_W_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_W_DXA}; got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_W_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_specs = {
        "Title": (24, BLUE_DARK, 0, 12),
        "Subtitle": (14, MUTED, 0, 8),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
        "Caption": (9, MUTED, 2, 10),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name in {"Title", "Heading 1", "Heading 2", "Heading 3"}
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")
        if name == "Caption":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Keep fields static. Automatic TOC refresh can make Word's PDF export
    # stall on some Windows installations; this manual uses a curated TOC.
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)


def add_field(paragraph, instruction: str, result: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate])
    if result:
        result_run = paragraph.add_run(result)
        set_run_font(result_run, 9, color=MUTED)
    paragraph.add_run()._r.append(end)


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [7200, 2160], indent=0)
    left, right = table.rows[0].cells
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("WMS | คู่มือการใช้งานระบบ"), 9.5, True, BLUE_DARK)
    p.add_run("\n")
    set_run_font(p.add_run("ฉบับภาษาไทย · ทุกบทบาท"), 8.5, color=MUTED)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("Rev. 1.0\nหน้า "), 8.5, color=MUTED)
    add_field(p, "PAGE", "1")
    for cell in (left, right):
        set_cell_border(cell, bottom={"val": "single", "sz": 6, "color": LIGHT_LINE})

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [1800, 7560], indent=0)
    logo_cell, company_cell = table.rows[0].cells
    logo_p = logo_cell.paragraphs[0]
    logo_p.paragraph_format.space_after = Pt(0)
    footer_logo = logo_p.add_run().add_picture(str(LOGO), width=Inches(0.72))
    footer_logo._inline.docPr.set("descr", "โลโก้บริษัท iCreativeSystems")
    company_p = company_cell.paragraphs[0]
    company_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company_p.paragraph_format.space_after = Pt(0)
    set_run_font(company_p.add_run("iCreativeSystems Co., Ltd.\nicsco.ai"), 8.5, color=MUTED)

    first_footer = section.first_page_footer
    p = first_footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("iCreativeSystems Co., Ltd.  ·  icsco.ai"), 9, color=MUTED)


def create_abstract_num(doc: Document, fmt: str, text: str, font_name: str | None = None) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, ppr])
    if font_name:
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)
        rpr.append(rfonts)
        lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def create_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = (max(num_ids) + 1) if num_ids else 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), str(abstract_id))
    num.append(abstract)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    ppr.append(num_pr)


def add_p(doc: Document, text: str = "", *, bold: bool = False, color: str = INK,
          size: float = 11, align=WD_ALIGN_PARAGRAPH.LEFT, before: float = 0,
          after: float = 6, keep_next: bool = False) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_with_next = keep_next
    run = p.add_run(text)
    set_run_font(run, size, bold, color)
    return p


def add_rich_p(doc: Document, segments: Iterable[tuple[str, bool]], *, after: float = 6) -> object:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    for text, bold in segments:
        set_run_font(p.add_run(text), 11, bold, INK)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, {1: 16, 2: 13, 3: 12}.get(level, 11), True,
                     BLUE if level < 3 else BLUE_DARK)
    return p


def add_list(doc: Document, items: Sequence[str], abstract_id: int, numbered: bool = False) -> None:
    num_id = create_num_instance(doc, abstract_id)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        apply_num(p, num_id)
        set_run_font(p.add_run(item), 10.5, color=INK)


def add_callout(doc: Document, title: str, text: str, tone: str = "info") -> None:
    fills = {"info": PALE_BLUE, "tip": PALE_TEAL, "warning": PALE_YELLOW,
             "danger": PALE_RED, "success": PALE_GREEN}
    accents = {"info": BLUE, "tip": TEAL, "warning": "D69E2E",
               "danger": "C53030", "success": "2F855A"}
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_W_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    set_cell_border(cell, left={"val": "single", "sz": 18, "color": accents[tone]},
                    top={"val": "single", "sz": 2, "color": fills[tone]},
                    bottom={"val": "single", "sz": 2, "color": fills[tone]},
                    right={"val": "single", "sz": 2, "color": fills[tone]})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title), 10.5, True, accents[tone])
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.2
    set_run_font(p.add_run(text), 10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]],
              widths: Sequence[int], alignments: Sequence | None = None) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.rows[0].height_rule = None
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, TABLE_HEADER)
        set_cell_border(cell, top={"color": LIGHT_LINE}, bottom={"color": LIGHT_LINE},
                        left={"color": LIGHT_LINE}, right={"color": LIGHT_LINE})
        p = cell.paragraphs[0]
        p.alignment = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), 9.5, True, BLUE_DARK)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_border(cells[idx], top={"color": LIGHT_LINE}, bottom={"color": LIGHT_LINE},
                            left={"color": LIGHT_LINE}, right={"color": LIGHT_LINE})
            p = cells[idx].paragraphs[0]
            p.alignment = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(str(value)), 9.2, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, filename: str, caption: str, *, max_width: float = 6.25,
               max_height: float = 6.7, page_break: bool = False) -> None:
    image_path = SHOT_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    if page_break:
        marker = doc.add_paragraph()
        marker.paragraph_format.page_break_before = True
        marker.paragraph_format.space_before = Pt(0)
        marker.paragraph_format.space_after = Pt(0)
        marker.paragraph_format.line_spacing = Pt(1)
        set_run_font(marker.add_run(""), 1, color=WHITE)
    with Image.open(image_path) as im:
        px_w, px_h = im.size
    width = min(max_width, max_height * px_w / px_h)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(width))
    inline._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(style="Caption")
    cap.paragraph_format.keep_with_next = False
    set_run_font(cap.add_run(caption), 9, color=MUTED)


def add_role_opener(doc: Document, role: str, thai_name: str, tagline: str,
                    can_do: Sequence[str], cannot_do: Sequence[str], accent: str) -> None:
    marker = doc.add_paragraph()
    marker.paragraph_format.page_break_before = True
    marker.paragraph_format.space_before = Pt(0)
    marker.paragraph_format.space_after = Pt(0)
    marker.paragraph_format.line_spacing = Pt(1)
    set_run_font(marker.add_run(""), 1, color=WHITE)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_W_DXA], indent=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, accent)
    set_cell_margins(cell, top=260, start=240, bottom=260, end=240)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(f"คู่มือสำหรับ {role}"), 22, True, WHITE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{thai_name} · {tagline}"), 12, color=WHITE)
    add_heading(doc, "ขอบเขตการใช้งาน", 1)
    add_rich_p(doc, [("ใช้บทนี้เมื่อ ", True),
                     (f"บัญชีของคุณแสดง Role เป็น {role} ในเมนูโปรไฟล์", False)])
    add_heading(doc, "ทำได้", 2)
    add_list(doc, can_do, BULLET_ABSTRACT)
    add_heading(doc, "ทำไม่ได้ / ต้องส่งต่อ", 2)
    add_list(doc, cannot_do, BULLET_ABSTRACT)


def add_cover(doc: Document) -> None:
    add_p(doc, "", before=78, after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_logo = p.add_run().add_picture(str(LOGO), width=Inches(3.05))
    cover_logo._inline.docPr.set("descr", "โลโก้บริษัท iCreativeSystems")
    p.paragraph_format.space_after = Pt(30)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("คู่มือการใช้งานระบบ WMS"), 24, True, BLUE_DARK)
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("User Guideline ฉบับภาษาไทย · สำหรับทุกบทบาท"), 14, color=MUTED)
    add_p(doc, "Viewer  ·  Operator  ·  Manager  ·  Admin",
          size=11.5, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=14, after=48)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_W_DXA], indent=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLUE_DARK)
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("ฉบับแก้ไข 1.0  |  22 กรกฎาคม 2569"), 10.5, True, WHITE)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run("เอกสารใช้งานภายใน · ตัวอย่างข้อมูลในภาพไม่มีผลต่อข้อมูลจริง"), 9, color=WHITE)
    doc.add_page_break()


def add_document_control(doc: Document) -> None:
    add_heading(doc, "ข้อมูลเอกสาร", 1)
    add_table(doc, ["รายการ", "รายละเอียด"], [
        ["ชื่อเอกสาร", "คู่มือการใช้งานระบบ WMS ฉบับภาษาไทย สำหรับทุกบทบาท"],
        ["เวอร์ชัน", "1.0"],
        ["วันที่จัดทำ", "22 กรกฎาคม 2569"],
        ["กลุ่มผู้ใช้", "Viewer, Operator, Manager และ Admin"],
        ["รูปแบบ", "Microsoft Word (.docx) สามารถแก้ไขข้อความและแทนที่ภาพได้"],
        ["การควบคุมเอกสาร", "ตรวจสอบให้ตรงกับเวอร์ชันระบบก่อนนำไปอบรมหรือประกาศใช้"],
    ], [2200, 7160])

    add_heading(doc, "วิธีใช้คู่มือนี้", 1)
    add_list(doc, [
        "เริ่มจากบท “การใช้งานร่วมกันทุกบทบาท” เพื่อสมัครสมาชิก เข้าสู่ระบบ และรู้จักเมนูหลัก",
        "ไปยังบทของ Role ที่ปรากฏในเมนูโปรไฟล์ของคุณ ไม่ควรใช้สิทธิ์หรือบัญชีของผู้อื่น",
        "ทำตามขั้นตอนหมายเลขตามลำดับ และอ่านข้อความ “จุดตรวจสอบ” ก่อนจบงานทุกครั้ง",
        "ภาพหน้าจอเป็นภาพจากระบบจริงโดยใช้ข้อมูลตัวอย่าง ตำแหน่งหรือจำนวนรายการอาจต่างจากระบบงานจริง",
        "คำว่า “ผู้ดูแล” ในเอกสาร หมายถึง Manager หรือ Admin เว้นแต่ระบุเป็นอย่างอื่น",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "หลักการสำคัญ", "ใช้บัญชีของตนเองเท่านั้น ระบบอนุญาตให้ 1 บัญชีใช้งานพร้อมกันได้ 1 อุปกรณ์ การเข้าสู่ระบบจากอุปกรณ์ใหม่จะทำให้เซสชันเดิมสิ้นสุด", "warning")

    add_heading(doc, "สัญลักษณ์ที่ใช้ในเอกสาร", 2)
    add_table(doc, ["ป้าย", "ความหมาย"], [
        ["ขั้นตอน", "สิ่งที่ผู้ใช้ต้องทำตามลำดับ"],
        ["จุดตรวจสอบ", "ผลที่ควรเห็นก่อนถือว่างานสำเร็จ"],
        ["เคล็ดลับ", "วิธีทำงานให้เร็วขึ้นหรือลดความผิดพลาด"],
        ["ข้อควรระวัง", "การกระทำที่มีผลต่อสต็อก ผู้ใช้ หรือข้อมูลถาวร"],
    ], [1800, 7560])
def add_toc(doc: Document) -> None:
    add_heading(doc, "สารบัญ", 1)
    add_p(doc, "โครงสร้างเอกสารจัดตามลำดับการเรียนรู้: เริ่มจากการใช้งานร่วมกัน แล้วเข้าสู่บทของแต่ละ Role", color=MUTED)
    add_table(doc, ["หัวข้อ", "เนื้อหาหลัก"], [
        ["1. ภาพรวมระบบ WMS", "Inbound, Outbound, Inventory, Report, Analysis และ User Admin"],
        ["2. บทบาทและสิทธิ์", "ตารางเปรียบเทียบ Viewer, Operator, Manager และ Admin"],
        ["3. การใช้งานร่วมกันทุกบทบาท", "สมัคร, Login, ลืมรหัสผ่าน, เมนู, แจ้งเตือน, สแกน, ตั้งค่า"],
        ["4. คู่มือ Viewer", "ดูแดชบอร์ด ตรวจสต็อก และสร้างรายงาน PDF"],
        ["5. คู่มือ Operator", "สร้างใบเบิก ติดตามผล ยกเลิก และรับสินค้า"],
        ["6. คู่มือ Manager", "สินค้า รับเข้า ปรับยอด CSV อนุมัติ ส่งมอบ และวิเคราะห์"],
        ["7. คู่มือ Admin", "อนุมัติผู้สมัคร เปลี่ยน Role ระงับ คืนสิทธิ์ และลบผู้ใช้"],
        ["8–10. รายงาน การแก้ปัญหา และความปลอดภัย", "เช็กลิสต์คุณภาพ FAQ ข้อมูลแจ้งปัญหา และการควบคุมงาน"],
    ], [3540, 5820])
    doc.add_page_break()


def add_system_overview(doc: Document) -> None:
    add_heading(doc, "1. ภาพรวมระบบ WMS", 1)
    add_p(doc, "WMS (Warehouse Management System) ใช้บันทึกและติดตามอะไหล่หรือสินค้า ตั้งแต่รับเข้าคลัง ขอเบิก พิจารณาใบเบิก ส่งมอบ ไปจนถึงรายงานและการวิเคราะห์สต็อก")
    add_table(doc, ["กระบวนการ", "สิ่งที่ระบบทำ", "ผู้รับผิดชอบหลัก"], [
        ["Inbound", "บันทึกรับสินค้าเข้าคลังและเพิ่มยอดคงเหลือ", "Manager / Admin"],
        ["Outbound", "Operator ส่งใบเบิก ผู้ดูแลพิจารณา แล้วส่งมอบ", "Operator + Manager / Admin"],
        ["Inventory", "ค้นหา กรอง สแกน และตรวจสอบยอดคงเหลือ", "ทุก Role"],
        ["Report", "สร้าง PDF ย้อนหลังรายวัน รายเดือน หรือรายปี", "ทุก Role"],
        ["Analysis", "ดูตัวชี้วัด แนวโน้ม และคาดการณ์สต็อก", "Manager / Admin"],
        ["User Admin", "อนุมัติ ระงับ เปลี่ยน Role และลบบัญชี", "Admin เท่านั้น"],
    ], [1800, 4960, 2600])

    add_heading(doc, "1.1 ลำดับงานเบิกสินค้า", 2)
    add_p(doc, "Operator เลือกสินค้า → ส่งใบเบิก → Manager/Admin ตรวจสอบ → ระบบตัดสต็อกเมื่ออนุมัติ → ผู้ขอมารับ → ผู้ดูแลกดส่งมอบแล้ว")
    add_callout(doc, "จุดที่สต็อกเปลี่ยน", "การส่งใบเบิกยังไม่ทำให้ยอดลด ระบบลดสต็อกเมื่อ Manager หรือ Admin บันทึกการอนุมัติเท่านั้น", "info")

    heading = add_heading(doc, "2. บทบาทและสิทธิ์", 1)
    heading.paragraph_format.page_break_before = True
    center = WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc, ["ความสามารถ", "Viewer", "Operator", "Manager", "Admin"], [
        ["ดูแดชบอร์ด / สินค้าคงคลัง", "ได้", "ได้", "ได้", "ได้"],
        ["ออกรายงาน PDF", "ได้", "ได้", "ได้", "ได้"],
        ["ขอเบิกสินค้า", "ไม่ได้", "ได้", "ได้", "ได้"],
        ["จัดการสินค้า / รับเข้า", "ไม่ได้", "ไม่ได้", "ได้", "ได้"],
        ["ปิดใช้ / คืนสถานะ / ลบสินค้า", "ไม่ได้", "ไม่ได้", "ได้", "ได้"],
        ["พิจารณา / ส่งมอบใบเบิก", "ไม่ได้", "ไม่ได้", "ได้", "ได้"],
        ["นำเข้า / ส่งออก CSV", "ไม่ได้", "ไม่ได้", "ได้", "ได้"],
        ["ดูหน้าวิเคราะห์", "ไม่ได้", "ไม่ได้", "ได้", "ได้"],
        ["จัดการผู้ใช้งาน", "ไม่ได้", "ไม่ได้", "ไม่ได้", "ได้"],
    ], [3320, 1400, 1500, 1540, 1600],
       [WD_ALIGN_PARAGRAPH.LEFT, center, center, center, center])
    add_callout(doc, "สรุป Role", "Viewer = ดูข้อมูล · Operator = ขอเบิกได้ · Manager = จัดการคลังได้ · Admin = จัดการผู้ใช้ได้เพิ่มจาก Manager", "tip")


def add_common_usage(doc: Document) -> None:
    heading = add_heading(doc, "3. การใช้งานร่วมกันทุกบทบาท", 1)
    heading.paragraph_format.page_break_before = True
    add_heading(doc, "3.1 สมัครสมาชิก", 2)
    add_list(doc, [
        "เปิดหน้าเข้าสู่ระบบ แล้วเลือก “สมัครสมาชิก”",
        "กรอกชื่อผู้ใช้ อีเมล รหัสผ่านอย่างน้อย 8 ตัวอักษร และยืนยันรหัสผ่านให้ตรงกัน",
        "กด “สมัครสมาชิก” ระบบจะสร้างบัญชีสถานะรออนุมัติและกำหนด Role เริ่มต้นเป็น Viewer",
        "รอ Admin อนุมัติ เมื่ออนุมัติแล้วจึงเข้าสู่ระบบได้ และ Admin สามารถปรับ Role ภายหลัง",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "register.png", "ภาพที่ 1 หน้าสมัครสมาชิกของระบบ WMS")
    add_callout(doc, "จุดตรวจสอบ", "หลังส่งแบบฟอร์มควรเห็นข้อความแจ้งว่าสมัครสำเร็จและให้รอการอนุมัติ หากยังเข้าไม่ได้ ให้ Admin ตรวจสอบสถานะ Pending", "success")

    add_heading(doc, "3.2 เข้าสู่ระบบ", 2)
    add_list(doc, [
        "กรอกชื่อผู้ใช้และรหัสผ่านของตนเอง",
        "เลือก “จดจำชื่อผู้ใช้” หากเป็นอุปกรณ์ส่วนตัว ระบบจะจำเฉพาะชื่อผู้ใช้ ไม่ใช่รหัสผ่าน",
        "กด “เข้าสู่ระบบ” แล้วรอให้หน้าแดชบอร์ดแสดงข้อมูล",
        "ตรวจสอบ Role ที่เมนูโปรไฟล์ก่อนเริ่มทำรายการ โดยเฉพาะงานที่เปลี่ยนสต็อกหรือจัดการผู้ใช้",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "login.png", "ภาพที่ 2 หน้าเข้าสู่ระบบ WMS")
    add_callout(doc, "เมื่อถูกตัดออกจากระบบ", "หากมีข้อความว่าบัญชีถูกเข้าสู่ระบบจากอุปกรณ์อื่น แสดงว่ามีการ Login บัญชีเดียวกันที่เครื่องใหม่ ให้ใช้บัญชีของตนเองหรือประสาน Admin", "warning")

    add_heading(doc, "3.3 ลืมรหัสผ่าน", 2)
    add_list(doc, [
        "ที่หน้าเข้าสู่ระบบ เลือก “ลืมรหัสผ่าน?”",
        "กรอกอีเมลที่ใช้สมัคร แล้วกด “ส่งลิงก์รีเซ็ตรหัสผ่าน”",
        "เปิดลิงก์จากอีเมลภายใน 30 นาที และตั้งรหัสผ่านใหม่อย่างน้อย 8 ตัวอักษร",
        "เข้าสู่ระบบด้วยรหัสผ่านใหม่ ลิงก์เดิมใช้ได้ครั้งเดียว หากหมดอายุให้ขอลิงก์ใหม่",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "forgot-password.png", "ภาพที่ 3 หน้าขอลิงก์รีเซ็ตรหัสผ่าน")

    doc.add_page_break()
    add_heading(doc, "3.4 ส่วนประกอบหลักหลังเข้าสู่ระบบ", 2)
    add_table(doc, ["ส่วน", "หน้าที่", "วิธีใช้อย่างย่อ"], [
        ["แถบเมนู", "ไปยังแดชบอร์ด สินค้าคงคลัง สินค้า วิเคราะห์ หรือผู้ใช้ตามสิทธิ์", "เลือกชื่อเมนู ระบบซ่อนเมนูที่ Role เข้าไม่ได้"],
        ["กระดิ่งแจ้งเตือน", "แจ้งใบเบิก ผลอนุมัติ หรือคำขอสมัครแบบเรียลไทม์", "เลือกข้อความเพื่อไปหน้าที่เกี่ยวข้อง หรือกดอ่าน/ล้างทั้งหมด"],
        ["สว่าง/มืด", "เปลี่ยนธีมหน้าจอ", "กดสัญลักษณ์ธีมบนแถบด้านบน ระบบจำค่าที่เลือก"],
        ["เมนูโปรไฟล์", "ตรวจ Role เปิดตั้งค่าบัญชี และออกจากระบบ", "เปลี่ยนชื่อ อีเมล รหัสผ่าน หรือรูปโปรไฟล์ได้"],
        ["ติดตั้งแอป", "ติดตั้ง WMS แบบ PWA บนเครื่องที่รองรับ", "เปิดผ่าน HTTPS แล้วเลือกแบนเนอร์ “ติดตั้ง WMS เป็นแอป”"],
    ], [1800, 3560, 4000])

    add_heading(doc, "3.5 การแจ้งเตือน", 2)
    add_list(doc, [
        "Manager/Admin ได้รับแจ้งเมื่อมีใบเบิกใหม่ และ Admin ได้รับแจ้งคำขอสมัครสมาชิก",
        "Operator ได้รับผลอนุมัติของใบเบิกตนเอง รวมถึงข้อความ “มารับสินค้าได้”",
        "เลือกข้อความแจ้งเตือนเพื่อไปยังรายการที่เกี่ยวข้อง และตรวจสอบสถานะจริงบนแดชบอร์ดอีกครั้ง",
        "การแจ้งเตือนของเบราว์เซอร์อาจต้องอนุญาตสิทธิ์ หากไม่เด้ง ให้ตรวจการตั้งค่าเว็บไซต์และเสียงของอุปกรณ์",
    ], BULLET_ABSTRACT)

    add_heading(doc, "3.6 สแกนบาร์โค้ดหรือ QR", 2)
    add_list(doc, [
        "มือถือ/แท็บเล็ต: กดปุ่มกล้อง อนุญาตใช้กล้อง และจัดรหัสให้อยู่ในกรอบ",
        "คอมพิวเตอร์: กดปุ่มกล้องเพื่อเตรียมช่องค้นหา แล้วใช้เครื่องสแกน USB ยิงรหัส",
        "รหัสที่อ่านได้จะถูกใส่ในช่องค้นหาอัตโนมัติ ให้ตรวจชื่อสินค้าและ SKU ก่อนทำรายการ",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "ข้อกำหนดของกล้อง", "การใช้กล้องผ่านเว็บต้องเปิดระบบด้วย HTTPS บนโดเมนที่ใช้งานจริง หากเปิดด้วย HTTP กล้องอาจไม่ทำงาน", "warning")

    add_heading(doc, "3.7 ตั้งค่าบัญชีและออกจากระบบ", 2)
    add_list(doc, [
        "เปิดเมนูโปรไฟล์มุมขวาบน แล้วเลือก “ตั้งค่าบัญชี”",
        "แก้ชื่อผู้ใช้ อีเมล รูปโปรไฟล์ หรือกรอกรหัสผ่านใหม่เมื่อต้องการเปลี่ยน",
        "บันทึกและตรวจข้อความยืนยัน จากนั้นกลับไปดูเมนูโปรไฟล์ว่าข้อมูลเปลี่ยนแล้ว",
        "เมื่อเลิกใช้งานให้เลือก “ออกจากระบบ” โดยเฉพาะเครื่องส่วนกลาง ไม่ควรปิดหน้าต่างทิ้งไว้เฉย ๆ",
    ], NUMBER_ABSTRACT)


def add_viewer_manual(doc: Document) -> None:
    add_role_opener(doc, "Viewer", "ผู้ชม", "ดูข้อมูลและออกรายงานแบบอ่านอย่างเดียว",
                    ["ดูแดชบอร์ดและประวัติการทำรายการ", "ค้นหา กรอง และสแกนสินค้าคงคลัง",
                     "ตรวจยอดคงเหลือ หมวดหมู่ และสถานะสต็อก", "สร้างรายงาน PDF ตามช่วงเวลา"],
                    ["ไม่มีปุ่มเพิ่มลงใบเบิก", "แก้ไขสินค้า รับเข้า หรือปรับยอดไม่ได้",
                     "อนุมัติใบเบิกและจัดการผู้ใช้ไม่ได้"], BLUE)
    add_heading(doc, "4.1 ตรวจภาพรวมบนแดชบอร์ด", 2)
    add_list(doc, [
        "เปิดเมนู “แดชบอร์ด”",
        "อ่านการ์ดจำนวนสินค้า สต็อกต่ำ รับเข้าวันนี้ และเบิกออกวันนี้",
        "เลื่อนดูประวัติการทำรายการและสถานะของแต่ละใบ",
        "หากตัวเลขผิดปกติ ให้จดช่วงเวลา SKU หรือเลขใบรายการ แล้วแจ้ง Manager/Admin ตรวจสอบ",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "จุดตรวจสอบ", "Viewer สามารถดูข้อมูลได้ แต่ไม่ควรเห็นปุ่มจัดการสินค้า ปุ่มพิจารณาใบเบิก หรือเมนูจัดการผู้ใช้", "info")

    add_heading(doc, "4.2 ค้นหาและตรวจสต็อก", 2)
    add_list(doc, [
        "เปิดเมนู “สินค้าคงคลัง”",
        "พิมพ์ชื่อสินค้า SKU หรือใช้ปุ่มกล้องเพื่อสแกน",
        "ใช้ตัวกรองหมวดหมู่เพื่อลดจำนวนรายการ",
        "ตรวจชื่อ SKU หน่วยนับ ยอดคงเหลือ และสถานะสต็อกให้ตรงกับสิ่งที่ต้องการ",
        "หากไม่พบสินค้า ให้ล้างคำค้น/ตัวกรองก่อน แล้วจึงแจ้งผู้ดูแลตรวจว่าสินค้าถูกปิดใช้งานหรือไม่",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "viewer-inventory.png", "ภาพที่ 4 หน้าสินค้าคงคลังสำหรับ Viewer — แสดงข้อมูลแบบอ่านอย่างเดียว")

    add_heading(doc, "4.3 สร้างรายงาน PDF", 2)
    add_list(doc, [
        "กลับไปหน้าแดชบอร์ด แล้วกด “นำออก PDF”",
        "เลือกประเภทรายงานรายวัน รายเดือน หรือรายปี และระบุวันที่/เดือน/ปี",
        "เลือกประเภทรายการทั้งหมด นำเข้า หรือเบิกออก และเลือกโปรเจกต์เมื่อใช้กับรายการเบิก",
        "ตรวจจำนวนข้อมูลที่ตรงเงื่อนไข แล้วกด “ดาวน์โหลดไฟล์ PDF”",
        "เปิดไฟล์ที่ดาวน์โหลด ตรวจหัวรายงาน ช่วงเวลา จำนวนรายการ และยอดรวมก่อนส่งต่อ",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "เมื่อปุ่มดาวน์โหลดใช้ไม่ได้", "ระบบจะปิดปุ่มเมื่อยังค้นหาอยู่หรือไม่พบข้อมูลตามเงื่อนไข ให้เปลี่ยนช่วงเวลา/ตัวกรองและรอจำนวนรายการอัปเดต", "tip")

    add_heading(doc, "4.4 เช็กลิสต์ก่อนจบงาน", 2)
    add_list(doc, ["ล้างคำค้นและตัวกรองที่ค้างอยู่", "ตรวจว่าไฟล์รายงานเปิดได้และช่วงเวลาถูกต้อง",
                   "ออกจากระบบเมื่อใช้เครื่องส่วนกลาง"], BULLET_ABSTRACT)


def add_operator_manual(doc: Document) -> None:
    add_role_opener(doc, "Operator", "พนักงาน", "ดูข้อมูลและส่งคำขอเบิกสินค้า",
                    ["ใช้งานทุกอย่างของ Viewer", "เพิ่มสินค้าได้หลายรายการลงใบเบิก",
                     "ระบุจำนวนและโปรเจกต์/เหตุผล", "ส่ง ติดตาม และยกเลิกใบเบิกที่ยัง Pending"],
                    ["แก้ยอดคงเหลือหรือรับสินค้าเข้าไม่ได้", "อนุมัติ/ปฏิเสธใบเบิกไม่ได้",
                     "จัดการสินค้า ผู้ใช้ หรือสิทธิ์ไม่ได้"], TEAL)
    add_heading(doc, "5.1 เตรียมก่อนสร้างใบเบิก", 2)
    add_list(doc, [
        "เตรียมชื่อโปรเจกต์หรือเหตุผลการใช้งาน และรายการสินค้า/จำนวนที่ต้องการ",
        "เปิดสินค้าคงคลัง ค้นหาด้วยชื่อ SKU หรือสแกนรหัส",
        "ตรวจว่าสินค้ามียอดมากกว่า 0 และชื่อ/หน่วยนับถูกต้อง",
        "หากสินค้าหายจากรายการ ให้แจ้ง Manager/Admin ตรวจสถานะปิดใช้งาน",
    ], NUMBER_ABSTRACT)

    add_heading(doc, "5.2 สร้างและส่งใบเบิก", 2)
    add_list(doc, [
        "กด “เพิ่มลงใบเบิก” ที่สินค้าที่ต้องการ ทำซ้ำได้หลายรายการ",
        "เปิด “ตะกร้าใบเบิก” มุมล่างขวา",
        "ปรับจำนวนด้วยปุ่มลบ/บวกหรือกรอกตัวเลข จำนวนต้องไม่เกินยอดที่ระบบอนุญาต",
        "ลบรายการที่ไม่ต้องการ และตรวจ SKU ชื่อ หน่วย และจำนวนทุกบรรทัด",
        "เลือกหรือกรอกชื่อโปรเจกต์/เหตุผลการเบิกให้ชัดเจน",
        "กด “ส่งใบเบิก” เพียงครั้งเดียว แล้วรอข้อความยืนยัน",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "operator-request-cart.png", "ภาพที่ 5 หน้าสรุปตะกร้าใบเบิกของ Operator")
    add_callout(doc, "จุดตรวจสอบ", "หลังส่งสำเร็จ ตะกร้าควรถูกล้าง และใบเบิกใหม่ควรปรากฏบนแดชบอร์ดในสถานะ “รออนุมัติ”", "success")

    add_heading(doc, "5.3 ติดตามผลและรับสินค้า", 2)
    add_table(doc, ["สถานะ", "ความหมาย", "สิ่งที่ Operator ต้องทำ"], [
        ["รออนุมัติ (Pending)", "ส่งใบเบิกแล้ว ผู้ดูแลยังไม่ตัดสิน", "รอ หรือยกเลิกหากยังไม่ต้องการ"],
        ["อนุมัติ", "อนุมัติครบและตัดสต็อกแล้ว", "อ่านข้อความและไปรับสินค้า"],
        ["อนุมัติบางส่วน", "จ่ายได้น้อยกว่าที่ขอ พร้อมเหตุผล", "ตรวจจำนวนที่อนุมัติแล้วไปรับ"],
        ["ปฏิเสธ", "ไม่อนุมัติใบเบิก พร้อมเหตุผล", "แก้ความต้องการหรือประสานผู้ดูแลก่อนส่งใหม่"],
        ["มารับสินค้าได้", "อนุมัติแล้วและรอส่งมอบ", "รับของที่คลัง ตรวจจำนวนต่อหน้าเจ้าหน้าที่"],
        ["ส่งมอบแล้ว", "ผู้ดูแลปิดขั้นตอนส่งมอบ", "เก็บเลขใบรายการเป็นหลักฐานหากจำเป็น"],
    ], [2100, 3460, 3800])
    add_list(doc, [
        "เปิดกระดิ่งแจ้งเตือนหรือหน้าแดชบอร์ดเพื่อตรวจผล",
        "เมื่อขึ้น “มารับสินค้าได้” ให้นำข้อมูลใบเบิกไปรับของที่คลัง",
        "ตรวจ SKU และจำนวนที่รับจริง หากไม่ตรง อย่าให้ปิดงานจนกว่าจะประสานผู้ดูแล",
        "ใบที่ยัง Pending สามารถยกเลิกได้ หลังอนุมัติแล้วต้องติดต่อผู้ดูแล",
    ], NUMBER_ABSTRACT)

    add_heading(doc, "5.4 ข้อผิดพลาดที่ควรหลีกเลี่ยง", 2)
    add_list(doc, [
        "กดส่งซ้ำเพราะหน้าเว็บตอบช้า — ให้รอข้อความยืนยันและตรวจแดชบอร์ดก่อน",
        "ใช้เหตุผลกว้างเกินไป เช่น “ใช้งาน” — ระบุชื่อโครงการ งาน หรือปลายทางให้ตรวจสอบย้อนหลังได้",
        "ขอจำนวนเกินความจำเป็น — อาจทำให้ผู้อื่นขาดของและทำให้ผู้ดูแลต้องอนุมัติบางส่วน",
        "ใช้บัญชีร่วมกัน — ทำให้ชื่อผู้ขอและประวัติไม่ตรงกับผู้รับของจริง",
    ], BULLET_ABSTRACT)


def add_manager_manual(doc: Document) -> None:
    add_role_opener(doc, "Manager", "ผู้จัดการคลัง", "จัดการสินค้า สต็อก ใบเบิก และการวิเคราะห์",
                    ["ใช้งานทุกอย่างของ Operator", "เพิ่ม แก้ไข รับเข้า ปรับยอด และปิด/คืนสถานะสินค้า",
                     "นำเข้า/ส่งออก CSV", "อนุมัติ ปฏิเสธ และยืนยันส่งมอบใบเบิก", "ดูการวิเคราะห์และคาดการณ์"],
                    ["อนุมัติหรือเปลี่ยน Role ผู้ใช้ไม่ได้", "ระงับหรือลบบัญชีผู้ใช้ไม่ได้",
                     "ไม่ควรลบสินค้าแบบถาวรโดยไม่มีการยืนยันผลกระทบ"], "355C7D")
    add_heading(doc, "6.1 ตรวจแดชบอร์ดต้นกะ", 2)
    add_list(doc, [
        "ตรวจการ์ดจำนวนสินค้า สต็อกต่ำ รับเข้า และเบิกออกวันนี้",
        "เปิดตารางคำขอเบิกรอดำเนินการ แยกใบ Pending ออกจากใบที่รอส่งมอบ",
        "ตรวจรายการเคลื่อนไหวล่าสุดเพื่อหายอดหรือกิจกรรมผิดปกติ",
        "กดการ์ดสต็อกต่ำเพื่อไปหน้าสินค้าที่กรองเฉพาะรายการต้องเติม",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "manager-dashboard.png", "ภาพที่ 6 แดชบอร์ดของ Manager พร้อมภาพรวมสต็อกและรายการดำเนินการ")

    doc.add_page_break()
    add_heading(doc, "6.2 เพิ่มสินค้าใหม่", 2)
    add_list(doc, [
        "เปิดเมนู “รายการอะไหล่” แล้วเลือกเพิ่มสินค้า",
        "เลือกหมวดหมู่ก่อน ระบบสามารถสร้าง SKU อัตโนมัติเมื่อเว้นเลขท้ายว่าง หรือกรอกเลขท้ายตามกติกาองค์กร",
        "กรอกชื่อสินค้า หน่วยนับ ผู้ขาย ราคาล่าสุด สต็อกขั้นต่ำ และรูปภาพตามข้อมูลที่ยืนยันแล้ว",
        "ตรวจว่า SKU ไม่ซ้ำและขึ้นต้นตามรหัสหมวด จากนั้นบันทึก",
        "ค้นหา SKU ที่เพิ่งสร้างเพื่อยืนยันชื่อ หมวด หน่วย และสถานะใช้งาน",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "ข้อควรระวังเรื่อง SKU", "SKU เป็นรหัสอ้างอิงหลักของสินค้า ตรวจหมวดและรูปแบบก่อนบันทึกหรือแก้ไข และหลีกเลี่ยงการเปลี่ยนรหัสหลังมีประวัติทำรายการโดยไม่วางแผนย้ายข้อมูล", "danger")

    add_heading(doc, "6.3 แก้ไข รับเข้า และปรับยอด", 2)
    add_table(doc, ["งาน", "ขั้นตอนย่อ", "จุดตรวจสอบ"], [
        ["แก้ไขสินค้า", "ค้นหาการ์ด → แก้ไข → เปลี่ยนข้อมูล → บันทึก", "ชื่อ หน่วย ผู้ขาย ขั้นต่ำ และรูปตรงกับเอกสารอ้างอิง"],
        ["รับอะไหล่เข้า", "กดรับเข้า → กรอกจำนวน/หมายเหตุ → ยืนยัน", "ยอดคงเหลือเพิ่มและมีรายการ Inbound ในประวัติ"],
        ["ปรับยอด", "เปิดการปรับยอด → ระบุยอด/สาเหตุ → ยืนยัน", "ยอดใหม่ตรงผลตรวจนับและเหตุผลตรวจสอบย้อนหลังได้"],
        ["สต็อกขั้นต่ำ", "แก้ค่า Min Stock ตามแผนเติมสินค้า", "สถานะสต็อกต่ำเปลี่ยนสัมพันธ์กับยอดจริง"],
    ], [1800, 4200, 3360])
    add_figure(doc, "admin-products.png", "ภาพที่ 7 หน้ารายการอะไหล่ที่ Manager/Admin ใช้จัดการสินค้า")

    add_heading(doc, "6.4 ปิดใช้งาน คืนสถานะ และลบถาวร", 2)
    add_list(doc, [
        "เลือก “ปิดใช้งาน” เมื่อต้องการซ่อนสินค้าออกจากหน้าค้นหา/เบิก แต่ยังเก็บประวัติไว้",
        "เปิดตัวกรอง “ที่ปิดใช้งาน” เพื่อดูรายการที่ถูกซ่อน และเลือกคืนสถานะเมื่อกลับมาใช้งาน",
        "ใช้ลบถาวรเฉพาะเมื่อยืนยันว่าไม่มีความจำเป็นต้องเก็บรายการและผลกระทบต่อประวัติได้รับการตรวจแล้ว",
        "หลังดำเนินการ ตรวจหน้าสินค้าคงคลังและรายงานว่ารายการแสดงตามที่ต้องการ",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "แนวปฏิบัติที่แนะนำ", "เลือกปิดใช้งานก่อนลบถาวรเสมอ เพราะคืนสถานะได้และยังรักษาความต่อเนื่องของข้อมูลย้อนหลัง", "warning")

    add_heading(doc, "6.5 นำเข้าและส่งออก CSV", 2)
    add_list(doc, [
        "ส่งออก CSV ปัจจุบันเพื่อใช้เป็นแม่แบบและสำรองข้อมูลก่อนแก้จำนวนมาก",
        "แก้เฉพาะคอลัมน์ที่ระบบรองรับ รักษาหัวคอลัมน์ รูปแบบ SKU และหน่วยนับให้สม่ำเสมอ",
        "นำเข้าไฟล์ แล้วอ่านผลสำเร็จ/ข้อผิดพลาดทุกบรรทัด ไม่ควรนำเข้าไฟล์เดิมซ้ำโดยไม่ตรวจผล",
        "สุ่มตรวจสินค้าหลายรายการหลังนำเข้า โดยเฉพาะ SKU หมวด ชื่อ สต็อกขั้นต่ำ และสถานะ",
    ], NUMBER_ABSTRACT)

    add_heading(doc, "6.6 พิจารณาใบเบิก", 2)
    add_list(doc, [
        "ที่แดชบอร์ด ในตารางคำขอเบิกรอดำเนินการ กด “ตรวจสอบ”",
        "ตรวจผู้ขอ โปรเจกต์ SKU ชื่อสินค้า จำนวนที่ขอ และยอดคงเหลือ",
        "กำหนดจำนวนอนุมัติแต่ละรายการ: เต็มจำนวน บางส่วน หรือ 0",
        "กรอกเหตุผลเมื่อปฏิเสธหรืออนุมัติไม่ครบ ระบบกำหนดให้ข้อมูลนี้จำเป็น",
        "กด “บันทึกการอนุมัติ” หรือ “ปฏิเสธทั้งใบ” เพียงครั้งเดียว",
        "ตรวจสถานะใหม่และยอดสต็อก ระบบตัดสต็อกตามจำนวนที่อนุมัติ",
    ], NUMBER_ABSTRACT)
    add_callout(doc, "ก่อนกดยืนยัน", "การอนุมัติมีผลต่อยอดสต็อกทันที ตรวจจำนวนอนุมัติทุกบรรทัดและข้อความถึงผู้ขอให้ครบถ้วน", "danger")

    add_heading(doc, "6.7 ส่งมอบสินค้า", 2)
    add_list(doc, [
        "เตรียมสินค้าตามจำนวนที่อนุมัติและเลขใบเบิก",
        "ตรวจตัวตนผู้รับ SKU หน่วย และจำนวนต่อหน้าผู้รับ",
        "เมื่อส่งครบแล้ว กด “ส่งมอบแล้ว ✓” เพื่อปิดใบเบิก",
        "หากส่งมอบไม่ครบหรือมีข้อโต้แย้ง อย่าปิดใบ ให้บันทึกและประสานผู้รับก่อน",
    ], NUMBER_ABSTRACT)

    add_heading(doc, "6.8 วิเคราะห์และวางแผนสต็อก", 2)
    add_list(doc, [
        "เปิดเมนู “วิเคราะห์” และเลือกช่วงเวลาหรือมุมมองที่ต้องการ",
        "ตรวจแนวโน้มรับเข้า/เบิกออก สินค้าเคลื่อนไหวสูง และรายการสต็อกต่ำ",
        "ใช้การคาดการณ์วันของหมดเป็นข้อมูลประกอบ ไม่ใช่คำสั่งซื้ออัตโนมัติ ต้องเทียบ Lead Time และแผนงานจริง",
        "บันทึก SKU และช่วงเวลาที่ใช้วิเคราะห์ เพื่อให้ผู้อื่นตรวจซ้ำได้",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "admin-analysis.png", "ภาพที่ 8 หน้าวิเคราะห์สต็อกและแนวโน้มสำหรับ Manager/Admin", max_width=5.1, max_height=7.25, page_break=True)


def add_admin_manual(doc: Document) -> None:
    add_role_opener(doc, "Admin", "ผู้ดูแลระบบ", "จัดการคลังครบถ้วนและควบคุมบัญชีผู้ใช้",
                    ["ใช้งานทุกอย่างของ Manager", "อนุมัติหรือปฏิเสธคำขอสมัคร",
                     "กำหนด Role และสถานะบัญชี", "กรอง ระงับ คืนสิทธิ์ และลบผู้ใช้"],
                    ["ลบบัญชีตนเองไม่ได้", "ลดสิทธิ์หรือลบ Admin คนสุดท้ายไม่ได้",
                     "ไม่ควรให้สิทธิ์เกินความจำเป็น"], "23395D")
    add_heading(doc, "7.1 ตรวจงานที่ต้องดูแล", 2)
    add_list(doc, [
        "ตรวจคำขอสมัครใหม่และใบเบิกค้างบนแดชบอร์ด",
        "ตรวจสต็อกต่ำ รายการเคลื่อนไหวผิดปกติ และการแจ้งเตือน",
        "เปิดหน้าจัดการผู้ใช้เพื่อตรวจสถานะ Pending/Denied และ Role ที่ไม่สอดคล้องกับหน้าที่",
        "ดำเนินงานคลังตามขั้นตอน Manager เมื่อจำเป็น",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "admin-dashboard.png", "ภาพที่ 9 แดชบอร์ดของ Admin")

    doc.add_page_break()
    add_heading(doc, "7.2 อนุมัติหรือปฏิเสธผู้สมัคร", 2)
    add_list(doc, [
        "เปิดเมนู “จัดการผู้ใช้งาน” และกรองสถานะ “รออนุมัติ”",
        "ตรวจชื่อผู้ใช้ อีเมล และยืนยันตัวบุคคลตามกระบวนการขององค์กร",
        "หากถูกต้อง กด “อนุมัติ” ระบบจะเปลี่ยนเป็น Active และส่งอีเมลแจ้งเมื่อระบบอีเมลพร้อมใช้งาน",
        "หากไม่ผ่าน ให้กด “ปฏิเสธ” และบันทึกเหตุผลตามกระบวนการภายใน",
        "กำหนด Role ให้ตรงหน้าที่จริง ผู้สมัครใหม่เริ่มต้นเป็น Viewer",
    ], NUMBER_ABSTRACT)
    add_figure(doc, "admin-users.png", "ภาพที่ 10 หน้าจัดการผู้ใช้งานสำหรับ Admin")
    add_callout(doc, "หลัก Least Privilege", "ให้สิทธิ์ต่ำสุดที่เพียงพอต่องาน: Viewer สำหรับดู, Operator สำหรับเบิก, Manager สำหรับจัดการคลัง และ Admin เฉพาะผู้ควบคุมบัญชี", "info")

    add_heading(doc, "7.3 เปลี่ยน Role", 2)
    add_list(doc, [
        "ค้นหาบัญชีด้วยชื่อหรืออีเมล และตรวจให้แน่ใจว่าเป็นบุคคลที่ต้องการ",
        "เลือก Role ใหม่จากรายการ Viewer, Operator, Manager หรือ Admin",
        "ยืนยันผลกระทบก่อนเพิ่มเป็น Manager/Admin เพราะสามารถเปลี่ยนสต็อกหรือสิทธิ์ได้",
        "แจ้งผู้ใช้ให้ออกจากระบบและเข้าสู่ระบบใหม่หากเมนูยังไม่อัปเดต",
        "บันทึกเหตุผลหรือผู้อนุมัติสิทธิ์ไว้ตามนโยบายองค์กร",
    ], NUMBER_ABSTRACT)

    add_heading(doc, "7.4 ระงับ คืนสิทธิ์ และลบผู้ใช้", 2)
    add_table(doc, ["การกระทำ", "ใช้เมื่อ", "ผลและการตรวจสอบ"], [
        ["ระงับ/ปฏิเสธ", "ไม่ควรให้บัญชีเข้าสู่ระบบชั่วคราว", "สถานะเป็น Denied และเข้าใช้งานไม่ได้"],
        ["คืนสิทธิ์", "ยืนยันแล้วว่าสามารถกลับมาใช้งาน", "สถานะ Active และ Role ยังถูกต้อง"],
        ["ลบผู้ใช้", "บัญชีไม่จำเป็นและผ่านการอนุมัติให้ลบ", "บัญชีถูกลบถาวร ต้องตรวจผลกระทบก่อน"],
    ], [2000, 3580, 3780])
    add_callout(doc, "ระบบป้องกันความเสียหาย", "ไม่สามารถลบบัญชีตนเอง และไม่สามารถลดสิทธิ์หรือลบ Admin คนสุดท้ายได้ ควรมี Admin ที่ได้รับมอบหมายอย่างน้อย 2 คนตามนโยบายองค์กร", "warning")

    add_heading(doc, "7.5 ทบทวนสิทธิ์เป็นระยะ", 2)
    add_list(doc, [
        "กรองผู้ใช้ตาม Active, Pending และ Denied แล้วตรวจบัญชีค้าง",
        "ทบทวน Manager/Admin ว่ายังมีหน้าที่จำเป็นและเป็นบุคคลปัจจุบัน",
        "ระงับบัญชีผู้ลาออก/ย้ายงานก่อนพิจารณาลบ เพื่อรักษาความสามารถในการตรวจสอบย้อนหลัง",
        "ห้ามใช้บัญชี Admin ร่วมกัน และไม่บันทึกรหัสผ่านในเอกสารคู่มือหรือไฟล์ส่วนกลาง",
    ], BULLET_ABSTRACT)


def add_reports_and_troubleshooting(doc: Document) -> None:
    # A fresh section keeps the long end-matter tables stable in Word's
    # pagination engine while preserving the running header and footer.
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    add_heading(doc, "8. รายงานและการตรวจสอบข้อมูล", 1)
    add_heading(doc, "8.1 เงื่อนไขรายงาน PDF", 2)
    add_table(doc, ["ตัวเลือก", "ค่าที่เลือกได้", "ข้อสังเกต"], [
        ["ประเภทรายงาน", "รายวัน / รายเดือน / รายปี", "กำหนดวันที่ เดือน หรือปีให้ตรงกับประเภท"],
        ["ประเภทรายการ", "ทั้งหมด / นำเข้า / เบิกออก", "ใช้กรองการเคลื่อนไหวที่ต้องการ"],
        ["โปรเจกต์", "ทุกโปรเจกต์หรือชื่อโปรเจกต์", "ใช้กับ Outbound; Inbound ไม่มีโปรเจกต์"],
        ["จำนวนรายการ", "ระบบแสดงก่อนดาวน์โหลด", "ถ้าเป็น 0 ให้เปลี่ยนช่วงเวลาหรือตัวกรอง"],
    ], [2000, 3140, 4220])
    add_heading(doc, "8.2 เช็กลิสต์คุณภาพรายงาน", 2)
    add_list(doc, [
        "หัวรายงานและช่วงเวลาตรงกับคำขอ",
        "ประเภท Inbound/Outbound และโปรเจกต์ตรงเงื่อนไข",
        "จำนวนรายการในไฟล์สอดคล้องกับจำนวนที่ระบบแจ้งก่อนดาวน์โหลด",
        "ชื่อผู้ทำรายการ สถานะ วันที่ และจำนวนอ่านได้ครบ",
        "ไฟล์เปิดได้ก่อนส่งต่อ และตั้งชื่อไฟล์ให้ค้นย้อนหลังได้",
    ], BULLET_ABSTRACT)

    doc.add_page_break()
    add_heading(doc, "9. แก้ปัญหาเบื้องต้น", 1)
    add_table(doc, ["อาการ", "สาเหตุที่พบบ่อย", "วิธีตรวจและแก้"], [
        ["สมัครแล้วเข้าไม่ได้", "บัญชียัง Pending", "ให้ Admin ตรวจและอนุมัติบัญชี"],
        ["ถูกเด้งออกจากระบบ", "บัญชีเดียวกัน Login ที่อุปกรณ์อื่น", "ใช้บัญชีส่วนตัวและเข้าสู่ระบบใหม่"],
        ["ลิงก์รีเซ็ตใช้ไม่ได้", "เกิน 30 นาทีหรือใช้ไปแล้ว", "ขอลิงก์ใหม่จากหน้าลืมรหัสผ่าน"],
        ["เบิกแล้วสต็อกไม่ลด", "ใบเบิกยังไม่ถูกอนุมัติ", "ตรวจสถานะ Pending; ยอดลดเมื่ออนุมัติ"],
        ["หาสินค้าไม่เจอ", "มีตัวกรองหรือสินค้าถูกปิดใช้", "ล้างตัวกรอง แล้วให้ผู้ดูแลตรวจรายการปิดใช้"],
        ["กล้องไม่ทำงาน", "เปิดผ่าน HTTP หรือไม่ได้อนุญาตกล้อง", "ใช้ HTTPS และตรวจสิทธิ์กล้องของเว็บไซต์"],
        ["รายงานดาวน์โหลดไม่ได้", "ไม่พบข้อมูลหรือยังประมวลผล", "เปลี่ยนช่วงเวลา/ตัวกรองและรอจำนวนรายการ"],
        ["เมนูไม่ตรง Role", "สิทธิ์เพิ่งเปลี่ยนหรือเซสชันเก่า", "ออกจากระบบ เข้าใหม่ และให้ Admin ตรวจ Role"],
    ], [2200, 2860, 4300])

    add_heading(doc, "9.1 ข้อมูลที่ควรแจ้งผู้ดูแลเมื่อพบปัญหา", 2)
    add_list(doc, [
        "ชื่อผู้ใช้และ Role (ไม่ส่งรหัสผ่าน)",
        "วันที่ เวลา และหน้าจอที่พบปัญหา",
        "SKU เลขใบเบิก หรือชื่อโปรเจกต์ที่เกี่ยวข้อง",
        "ข้อความผิดพลาดที่เห็น และขั้นตอนก่อนเกิดปัญหา",
        "ภาพหน้าจอที่ไม่เปิดเผยข้อมูลลับหรือข้อมูลส่วนบุคคลเกินจำเป็น",
    ], BULLET_ABSTRACT)

    doc.add_page_break()
    add_heading(doc, "10. ความปลอดภัยและการควบคุมงาน", 1)
    add_list(doc, [
        "ใช้บัญชีของตนเองและออกจากระบบเมื่อเลิกใช้งาน โดยเฉพาะอุปกรณ์ส่วนกลาง",
        "ไม่ส่งรหัสผ่าน ลิงก์รีเซ็ต หรือข้อมูลเข้าสู่ระบบผ่านช่องทางที่ไม่ปลอดภัย",
        "ตรวจจำนวนก่อนอนุมัติ รับเข้า ปรับยอด และส่งมอบ เพราะมีผลต่อสต็อกและรายงาน",
        "ใช้การปิดใช้งานแทนลบถาวรเมื่อยังต้องเก็บประวัติ",
        "Admin ทบทวนสิทธิ์เป็นระยะและให้ Role ตามหน้าที่จริง",
        "เมื่อสงสัยว่าบัญชีถูกใช้โดยผู้อื่น ให้เปลี่ยนรหัสผ่านและแจ้ง Admin ทันที",
    ], BULLET_ABSTRACT)

    add_callout(doc, "สิ้นสุดคู่มือ", "ก่อนนำเอกสารไปใช้อบรม ให้เจ้าของระบบตรวจชื่อเมนู ขั้นตอนอนุมัติ และนโยบายองค์กรอีกครั้ง หาก UI เปลี่ยน ให้แทนที่ภาพพร้อมปรับคำอธิบายและเลข Revision", "tip")


def build(output_file: Path = OUT_FILE, stop_after: str = "all") -> Path:
    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "คู่มือการใช้งานระบบ WMS ฉบับภาษาไทย สำหรับทุกบทบาท"
    doc.core_properties.subject = "User guideline for Viewer, Operator, Manager, and Admin"
    doc.core_properties.author = "iCreativeSystems Co., Ltd."
    doc.core_properties.comments = "Generated from repository behavior and real WMS screenshots."

    global BULLET_ABSTRACT, NUMBER_ABSTRACT
    BULLET_ABSTRACT = create_abstract_num(doc, "bullet", "•", FONT)
    NUMBER_ABSTRACT = create_abstract_num(doc, "decimal", "%1.", FONT)

    builders = [
        ("cover", add_cover),
        ("control", add_document_control),
        ("toc", add_toc),
        ("overview", add_system_overview),
        ("common", add_common_usage),
        ("viewer", add_viewer_manual),
        ("operator", add_operator_manual),
        ("manager", add_manager_manual),
        ("admin", add_admin_manual),
        ("all", add_reports_and_troubleshooting),
    ]
    if stop_after == "reports_only":
        builders = [("reports_only", add_reports_and_troubleshooting)]
    valid_stops = {name for name, _ in builders}
    if stop_after not in valid_stops:
        raise ValueError(f"Unknown stop point: {stop_after}")
    for name, builder in builders:
        builder(doc)
        if name == stop_after:
            break

    doc.save(output_file)
    return output_file


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUT_FILE)
    parser.add_argument("--stop-after", default="all",
                        choices=["cover", "control", "toc", "overview", "common",
                                 "viewer", "operator", "manager", "admin", "all", "reports_only"])
    args = parser.parse_args()
    output = build(args.output, args.stop_after)
    print(output)
